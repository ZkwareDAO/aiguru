import streamlit as st
import os
import json
import hashlib
from datetime import datetime, timedelta
import time
import logging
from pathlib import Path
from functions.api_correcting.pdf_merger import PDFMerger, ImageToPDFConverter
from functions.api_correcting.calling_api import call_api, correction_with_image_marking_scheme
import re

# Constants
MAX_FILE_SIZE = 5 * 1024  # 5MB in KB
UPLOAD_DIR = Path("uploads")
DATA_FILE = Path("user_data.json")
MAX_RECORD_AGE_DAYS = 7  # Records older than this will be removed

# Create necessary directories
UPLOAD_DIR.mkdir(exist_ok=True)

# Test accounts for development
TEST_ACCOUNTS = {
    "test_user_1": {"password": "password1"},
    "test_user_2": {"password": "password2"}
}

try:
    from fpdf import FPDF
except ImportError:
    import subprocess
    import sys
    logging.info("Installing required package: fpdf")
    subprocess.check_call([sys.executable, "-m", "pip",
                          "install", "fpdf", "Pillow"])
    from fpdf import FPDF


def setup_logger(log_dir="logs"):
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)
    log_file = os.path.join(log_dir, "app_debug.log")
    logging.basicConfig(
        level=logging.INFO,
        format="[%(asctime)s] %(levelname)s @ %(module)s: %(message)s",
        handlers=[
            logging.FileHandler(log_file),
            logging.StreamHandler()
        ]
    )


# Initialize logger
setup_logger()
logging.info("Starting")

# Initialize storage structure
if not os.path.exists(DATA_FILE):
    with open(DATA_FILE, "w") as f:
        json.dump({}, f)


def filter_old_records(user_data):
    """
    Filter out records older than MAX_RECORD_AGE_DAYS days from user data

    Args:
        user_data (dict): The user data dictionary containing records

    Returns:
        dict: Updated user data with old records removed
    """
    current_date = datetime.now()
    cutoff_date = current_date - timedelta(days=MAX_RECORD_AGE_DAYS)
    records_removed = 0

    # Process each user's records
    for username, user_info in user_data.items():
        # Check if user_info is a dictionary and contains 'records' key
        if isinstance(
                user_info,
                dict) and 'records' in user_info and isinstance(
                user_info['records'],
                list):
            # Filter out old records
            filtered_records = []
            for record in user_info['records']:
                # Get record timestamp
                timestamp_str = record.get('timestamp')
                if not timestamp_str:
                    # Keep records without timestamp (shouldn't happen but just
                    # in case)
                    filtered_records.append(record)
                    continue

                try:
                    # Parse the timestamp
                    record_date = datetime.strptime(
                        timestamp_str, "%Y-%m-%d %H:%M:%S")
                    # Keep the record if it's newer than the cutoff date
                    if record_date >= cutoff_date:
                        filtered_records.append(record)
                    else:
                        records_removed += 1
                except (ValueError, TypeError):
                    # Keep records with invalid timestamps
                    filtered_records.append(record)

            # Update the user's records
            user_info['records'] = filtered_records

    if records_removed > 0:
        logging.info(
            f"Removed {records_removed} records older than {MAX_RECORD_AGE_DAYS} days")

    return user_data


def read_user_data():
    """从JSON文件读取用户数据，或返回默认数据"""
    try:
        with open(DATA_FILE, "r") as f:
            data = json.load(f)

            # 确保测试账户存在并使用哈希密码
            for test_user, details in TEST_ACCOUNTS.items():
                if test_user not in data:
                    data[test_user] = {
                        "password": details["password"],  # 对于测试账户，保持原始密码
                        "email": f"{test_user}@example.com",
                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "records": []
                    }

            # Filter out records older than MAX_RECORD_AGE_DAYS days
            data = filter_old_records(data)

            return data
    except FileNotFoundError:
        # 返回带有测试账户的默认数据
        default_data = {}
        for test_user, details in TEST_ACCOUNTS.items():
            default_data[test_user] = {
                "password": details["password"],
                "email": f"{test_user}@example.com",
                "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "records": []
            }
        return default_data


def save_user_data(data):
    with open(DATA_FILE, "w") as f:
        json.dump(data, f, indent=2)

# 添加辅助函数确保文件路径有效


def ensure_valid_file_path(file_path):
    """
    确保文件路径是有效的字符串并检查文件是否存在

    参数:
    file_path: 文件路径，可能是Path对象或字符串

    返回:
    tuple: (是否有效, 字符串形式的路径)
    """
    # 确保路径是字符串
    str_path = str(file_path) if file_path is not None else ""

    # 检查文件是否存在
    is_valid = os.path.exists(str_path) if str_path else False

    if not is_valid and str_path:
        logging.warning(f"File not found: {str_path}")

    return is_valid, str_path


def create_download_options(content, prefix="correction_result", files=None):
    """Create unified download options for content in TXT, PDF and Word formats

    Args:
        content: 文本内容
        prefix: 文件名前缀
        files: 字典，包含上传的图片文件信息 {'question': {}, 'student_answer': {}, 'marking_scheme': {}}
    """
    # 使用稳定的键名，不依赖于时间戳
    format_key = f"download_format_{prefix}"

    # 为了下载按钮创建一个唯一的时间戳
    current_time = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Create the dropdown menu for download options
    download_option = st.selectbox("Download Format",
                                   ["Select format...",
                                    "Text file (.txt)",
                                    "PDF Document (.pdf)",
                                    "Word Document (.docx)"],
                                   key=format_key)

    # Handle download based on selected option
    if download_option == "Text file (.txt)":
        st.download_button(
            label="Download result",
            data=content.encode('utf-8'),
            file_name=f"{prefix}_{current_time}.txt",
            mime="text/plain",
            key=f"download_txt_{prefix}_{current_time}"
        )

    elif download_option == "PDF Document (.pdf)":
        # 使用按钮触发 PDF 生成，避免自动生成导致问题
        if st.button(
            "Generate and Download PDF",
            key=f"generate_pdf_{prefix}_{current_time}",
                type="primary"):
            # 创建一个进度转轮和状态消息的容器
            progress_container = st.empty()
            status_container = st.empty()

            # 显示初始进度转轮
            with progress_container:
                progress_spinner = st.spinner("Initializing PDF generation...")

            # 显示初始状态消息
            with status_container:
                st.info("Starting PDF generation process...")

            with progress_spinner:
                try:
                    # 更新状态消息
                    status_container.info("Creating PDF merger instance...")

                    # Create PDFMerger instance
                    pdf_merger = PDFMerger(UPLOAD_DIR)

                    # 更新状态消息
                    status_container.info("Setting up output paths...")

                    # Generate PDF and provide download button
                    pdf_filename = f"{prefix}_{current_time}.pdf"
                    output_path = UPLOAD_DIR / st.session_state.current_user / pdf_filename

                    # 确保用户目录存在
                    user_dir = UPLOAD_DIR / st.session_state.current_user
                    user_dir.mkdir(exist_ok=True)

                    # 更新状态消息
                    status_container.info("Processing input files...")

                    # 准备文件供PDF合并器使用
                    files_to_include = {}

                    # 添加日志以诊断文件内容
                    logging.info(
                        f"Files dictionary for {prefix}: {
                            str(files)}")

                    if files and isinstance(files, dict):
                        # 更新状态消息
                        status_container.info("Validating uploaded files...")

                        # 遍历文件字典
                        for key, file_info in files.items():
                            if isinstance(
                                    file_info, dict) and 'saved_path' in file_info:
                                is_valid, file_path = ensure_valid_file_path(
                                    file_info['saved_path'])
                                if is_valid:
                                    # 根据文件类型确定文件标识符
                                    file_type = key

                                    files_to_include[file_type] = {
                                        'path': file_path}
                                    logging.info(
                                        f"Added file for PDF: {file_type} = {file_path}")
                                else:
                                    logging.warning(
                                        f"File does not exist: {
                                            file_info['saved_path']}")
                            else:
                                logging.warning(
                                    f"Invalid file info for {key}: {file_info}")
                    else:
                        logging.warning(
                            f"No valid files provided for {prefix} PDF generation")

                    logging.info(
                        f"Preparing to generate PDF with {
                            len(files_to_include)} files")

                    # 更新状态消息
                    status_container.info(
                        f"Generating PDF with {
                            len(files_to_include)} files...")

                    success, pdf_path = pdf_merger.merge_pdfs(
                        files_to_include,
                        content,
                        "AI Correction Results",
                        output_path
                    )

                    if success and os.path.exists(pdf_path):
                        # 更新状态消息
                        status_container.success("PDF generated successfully!")

                        logging.info(
                            f"PDF generated successfully at {pdf_path}")
                        with open(pdf_path, 'rb') as pdf_file:
                            pdf_bytes = pdf_file.read()

                            # 直接显示下载按钮
                            st.download_button(
                                label="Download PDF",
                                data=pdf_bytes,
                                file_name=pdf_filename,
                                mime="application/pdf",
                                key=f"download_pdf_{prefix}_{current_time}"
                            )

                        # Clean up temporary file
                        try:
                            if os.path.exists(pdf_path):
                                os.remove(pdf_path)
                                logging.info(
                                    f"Temporary PDF file removed: {pdf_path}")
                        except Exception as e:
                            logging.error(
                                f"Error removing temporary PDF file: {
                                    str(e)}")
                    else:
                        error_reason = "PDF file not found after generation" if not os.path.exists(
                            pdf_path) else "PDF merge operation failed"
                        status_container.error(
                            f"PDF generation failed: {error_reason}. Please try text format instead.")

                        # Show alternative download option
                        st.info("You can try downloading as text file instead.")
                        st.download_button(
                            label="Download as Text",
                            data=content.encode('utf-8'),
                            file_name=f"{prefix}_{current_time}.txt",
                            mime="text/plain",
                            key=f"fallback_txt_{prefix}_{current_time}"
                        )

                        if pdf_path:
                            logging.error(
                                f"PDF generation failed. Path: {pdf_path}, Success flag: {success}")
                        else:
                            logging.error(
                                "PDF generation failed. No path returned.")

                except Exception as e:
                    status_container.error(f"PDF generation failed: {str(e)}")
                    logging.error(
                        f"PDF generation error: {
                            str(e)}", exc_info=True)

                    # Show alternative download option after exception
                    st.info("You can try downloading as text file instead.")
                    st.download_button(
                        label="Download as Text",
                        data=content.encode('utf-8'),
                        file_name=f"{prefix}_{current_time}.txt",
                        mime="text/plain",
                        key=f"exception_txt_{prefix}_{current_time}"
                    )

    elif download_option == "Word Document (.docx)":
        try:
            # Import python-docx for Word document creation
            try:
                import docx
            except ImportError:
                import subprocess
                import sys
                logging.info("Installing required package: python-docx")
                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", "python-docx"])
                import docx

            from docx import Document
            import tempfile

            # 创建Word文档
            doc = Document()
            doc.add_heading("AI Correction Results", 0)

            # 添加内容
            for para in content.split('\n'):
                if para.strip():
                    doc.add_paragraph(para)

            # 保存到临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                docx_path = tmp.name
                doc.save(docx_path)

                # 读取文件内容用于下载
                with open(docx_path, 'rb') as f:
                    docx_data = f.read()

                # 提供下载按钮
                st.download_button(
                    label="Download Word Document",
                    data=docx_data,
                    file_name=f"{prefix}_{current_time}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_docx_{prefix}_{current_time}")

                # 清理临时文件
                try:
                    os.remove(docx_path)
                except BaseException:
                    pass

        except Exception as e:
            st.error(f"Word document generation failed: {str(e)}")
            logging.error(
                f"Word document generation error: {
                    str(e)}", exc_info=True)


def file_management_page():
    """File management center with navigation to history records"""
    st.title("📁 File Management Center")

    if not st.session_state.logged_in:
        st.warning("⚠️ Please log in to access file management features")
        if st.button("Log in now", type="primary"):
            st.session_state.show_login_modal = True
            st.session_state.next_page_after_login = "file_management"
            st.session_state.needs_rerun = True
        return

    # Card-style dashboard layout
    st.markdown("""
    <style>
    .dashboard-card {
        background-color: #f8f9fa;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        border-left: 5px solid #5b9bd5;
        transition: transform 0.3s;
    }
    .dashboard-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 20px rgba(0,0,0,0.1);
    }
    </style>
    """, unsafe_allow_html=True)

    # History Records Card
    st.markdown("""
    <div class="dashboard-card">
        <h3 style="color: #333;">📋 Correction History</h3>
        <p style="color: #666;">View and manage your previous correction records.</p>
    </div>
    """, unsafe_allow_html=True)

    if st.button(
        "View History Records",
        type="primary",
            use_container_width=True):
        st.session_state.page = "history_records"
        st.session_state.needs_rerun = True


def history_records_page():
    """Dedicated page for viewing and managing history records"""
    st.title("📋 Correction History Records")

    if not st.session_state.logged_in:
        st.warning("⚠️ Please log in to view your history records")
        if st.button("Log in now", type="primary"):
            st.session_state.show_login_modal = True
            st.session_state.next_page_after_login = "history_records"
            st.session_state.needs_rerun = True
        return

    # Navigation button to return to file management
    if st.button("← Back to File Management", key="back_to_file_mgmt"):
        st.session_state.page = "file_management"
        st.session_state.needs_rerun = True

    # Information about the 7-day record retention policy
    st.info(
        f"ℹ️ Records are automatically deleted after {MAX_RECORD_AGE_DAYS} days.")

    # Retrieve user data
    user_data = read_user_data()
    # Filter out records with empty content
    user_records = [
        record for record in user_data.get(
            st.session_state.current_user,
            {}).get(
            'records',
            []) if record.get('content') and record['content'].strip()]

    if not user_records:
        st.info("No correction records found.")
        return

    # Clear history functionality
    with st.expander("Clear History", expanded=False):
        st.warning(
            "This will permanently delete your correction history records.")

        # Only provide clear all option
        if st.button("Clear All Records", key="clear_all_btn", type="primary"):
            # Clear all records
            user_data[st.session_state.current_user]["records"] = []
            save_user_data(user_data)
            st.success("All records have been cleared!")
            st.session_state.needs_rerun = True

    # Display records in a grid of cards
    st.subheader("Your Correction Records")

    # Create a grid layout for the record cards
    num_columns = 3
    columns = st.columns(num_columns)

    # Custom CSS for cards
    st.markdown("""
    <style>
    .record-card {
        background-color: white;
        border-radius: 8px;
        padding: 0;
        margin-bottom: 15px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        height: 100%;
        transition: transform 0.2s, box-shadow 0.2s;
        overflow: hidden;
        cursor: pointer;
        position: relative;
    }
    .record-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
        background-color: #f8f9fa;
    }
    .record-card .invisible-button {
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        opacity: 0;
        z-index: 1;
        cursor: pointer;
    }
    .record-header {
        background-color: #f8f9fa;
        padding: 10px 15px;
        border-bottom: 1px solid #eee;
    }
    .record-title {
        font-weight: bold;
        font-size: 0.95em;
        color: #333;
        margin: 0;
    }
    .record-content {
        padding: 12px 15px;
    }
    .record-preview {
        color: #666;
        font-size: 0.85em;
        display: -webkit-box;
        -webkit-line-clamp: 3;
        -webkit-box-orient: vertical;
        overflow: hidden;
        line-height: 1.4;
        margin-bottom: 10px;
    }
    .card-footer {
        text-align: center;
        padding: 8px 0;
        background-color: #f0f8ff;
        border-top: 1px solid #eee;
    }
    .view-details-button {
        background-color: #5b9bd5;
        color: white;
        border: none;
        border-radius: 4px;
        padding: 4px 12px;
        font-size: 0.85em;
        font-weight: 500;
        cursor: pointer;
        transition: background-color 0.2s;
    }
    .view-details-button:hover {
        background-color: #4a8bc5;
    }
    </style>
    """, unsafe_allow_html=True)

    # Simplified JavaScript - only used for logging now
    st.markdown("""
    <script>
    console.log("Record cards JavaScript loaded");
    document.addEventListener('DOMContentLoaded', function() {
        console.log("DOM fully loaded");

        // Handle form submissions for navigation
        const recordForms = document.querySelectorAll('.record-nav-form');
        recordForms.forEach(form => {
            form.addEventListener('submit', function(e) {
                console.log('Form submitted:', this.id);
            });
        });
    });
    </script>
    """, unsafe_allow_html=True)

    # Add a session state dictionary to track clicks
    if 'record_clicks' not in st.session_state:
        st.session_state.record_clicks = {}

    # Check if any record was clicked via form submission
    query_params = st.query_params
    if 'view_record' in query_params:
        record_id = query_params['view_record']
        st.session_state.record_clicks[record_id] = True
        # Clear the query parameter to avoid repeated triggering
        query_params.clear()

    # Determine number of records per page for pagination
    RECORDS_PER_PAGE = 9

    # Add pagination
    if 'history_page' not in st.session_state:
        st.session_state.history_page = 0

    max_pages = len(user_records) // RECORDS_PER_PAGE
    if len(user_records) % RECORDS_PER_PAGE > 0:
        max_pages += 1

    # Pagination navigation
    if max_pages > 1:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            page_col1, page_col2, page_col3 = st.columns([1, 2, 1])
            with page_col1:
                if st.button(
                    "◀ Previous", disabled=(
                        st.session_state.history_page == 0)):
                    st.session_state.history_page -= 1
                    st.session_state.needs_rerun = True
            with page_col2:
                st.write(
                    f"Page {
                        st.session_state.history_page +
                        1} of {max_pages}")
            with page_col3:
                if st.button(
                    "Next ▶", disabled=(
                        st.session_state.history_page >= max_pages - 1)):
                    st.session_state.history_page += 1
                    st.session_state.needs_rerun = True

    # Display records for the current page
    start_idx = st.session_state.history_page * RECORDS_PER_PAGE
    end_idx = min(start_idx + RECORDS_PER_PAGE, len(user_records))
    current_page_records = user_records[start_idx:end_idx]

    for i, record in enumerate(current_page_records):
        col_idx = i % num_columns

        with columns[col_idx]:
            # Get record metadata
            timestamp = record.get('timestamp', 'No timestamp')
            settings = record.get('settings', {})
            if not settings or not isinstance(settings, dict):
                settings = {}

            strictness = settings.get('strictness_level', '中等')
            language = settings.get('language', 'zh')
            language_display = "Chinese (中文)" if language == "zh" else "English"

            # Record index
            record_index = len(user_records) - (start_idx + i)

            # Content preview
            content = record.get('content', 'No content available')

            # Extract summary information from content
            preview_lines = []
            for line in content.split('\n'):
                line = line.strip()
                # Look for key information like titles, score, subject
                if line.startswith('# ') or line.startswith(
                        '题目信息') or '总分' in line or '科目' in line:
                    preview_lines.append(line.replace('# ', ''))
                    if len(preview_lines) >= 3:
                        break

            # If we couldn't find structured information, fallback to a regular
            # preview
            if not preview_lines:
                preview = content[:150] + \
                    "..." if len(content) > 150 else content
            else:
                preview = "\n".join(preview_lines)

            # Create a regular button that's outside the HTML
            button_id = f"card_button_{record_index}"

            # Check if this record's button was clicked
            if st.button(
                    f"View Record {record_index}",
                    key=button_id,
                    use_container_width=True):
                # This direct button will handle the navigation properly
                st.session_state.selected_record_index = record_index
                st.session_state.selected_record = record
                st.session_state.page = "record_detail"
                st.session_state.needs_rerun = True
                logging.info(
                    f"Regular button clicked for record {record_index}")

            # Card container with simple HTML - styled to look like our cards
            # but using a direct button
            html_card = f"""
            <div class="record-card">
                <div class="record-header">
                    <div class="record-title">Record {record_index}: {timestamp}</div>
                </div>
                <div class="record-content">
                    <div class="record-preview">{preview}</div>
                </div>
                <div class="card-footer">
                </div>
            </div>
            """

            # Render the HTML card structure
            st.markdown(html_card, unsafe_allow_html=True)


def record_detail_page():
    """Detailed view of a specific correction record"""
    if 'selected_record' not in st.session_state:
        st.error("No record selected. Please go back to history records.")
        if st.button("Go to History Records"):
            st.session_state.page = "history_records"
            st.session_state.needs_rerun = True
        return

    record = st.session_state.selected_record
    record_index = st.session_state.selected_record_index

    # Navigation buttons
    col1, col2 = st.columns([1, 3])
    with col1:
        if st.button("← Back to Records", key="back_to_records"):
            st.session_state.page = "history_records"
            st.session_state.needs_rerun = True

    # Record metadata
    timestamp = record.get('timestamp', 'No timestamp')
    settings = record.get('settings', {})
    if not settings or not isinstance(settings, dict):
        settings = {}

    strictness = settings.get('strictness_level', '中等')
    language = settings.get('language', 'zh')
    language_display = "Chinese (中文)" if language == "zh" else "English"

    # Record header with integrated metadata - more compact design
    st.markdown("""
    <style>
    .record-detail-header {
        background-color: #f8f9fa;
        padding: 15px 20px;
        border-radius: 8px;
        margin-bottom: 20px;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
    }
    .record-detail-title {
        font-size: 1.5em;
        font-weight: bold;
        color: #333;
        margin-bottom: 10px;
        border-bottom: 1px solid #eee;
        padding-bottom: 10px;
    }
    .metadata-row {
        display: flex;
        flex-wrap: wrap;
        gap: 20px;
    }
    .metadata-item {
        display: inline-flex;
        align-items: center;
        color: #555;
    }
    .metadata-icon {
        margin-right: 5px;
        color: #5b9bd5;
    }
    .metadata-label {
        font-weight: bold;
        color: #333;
        margin-right: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="record-detail-header">
        <div class="record-detail-title">Record #{record_index}</div>
        <div class="metadata-row">
            <div class="metadata-item">
                <span class="metadata-icon">📅</span>
                <span class="metadata-label">Date:</span> {timestamp}
            </div>
            <div class="metadata-item">
                <span class="metadata-icon">🌐</span>
                <span class="metadata-label">Language:</span> {language_display}
            </div>
            <div class="metadata-item">
                <span class="metadata-icon">⚖️</span>
                <span class="metadata-label">Strictness:</span> {strictness}
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Buttons row - Delete and Download options
    col1, col2 = st.columns(2)
    with col1:
        # Delete record functionality
        if st.button(
            "🗑️ Delete This Record",
            type="secondary",
                use_container_width=True):
            user_data = read_user_data()
            user_records = user_data.get(
                st.session_state.current_user, {}).get(
                'records', [])

            # Find the record by timestamp to delete it
            record_timestamp = record.get('timestamp')
            for i, r in enumerate(user_records):
                if r.get('timestamp') == record_timestamp:
                    user_records.pop(i)
                    user_data[st.session_state.current_user]['records'] = user_records
                    save_user_data(user_data)
                    st.success("Record deleted successfully!")

                    # Return to history page
                    st.session_state.page = "history_records"
                    st.session_state.needs_rerun = True
                    break

    # 获取内容和文件信息
    content = record.get('content', 'No content available')
    files_dict = record.get('files', {})

    # 应用格式化来提高可读性
    st.markdown("""
    <style>
    .result-container {
        background-color: white;
        border-radius: 8px;
        padding: 20px;
        border-left: 5px solid #5cb85c;
        font-size: 0.95em;
        line-height: 1.5;
        box-shadow: 0 2px 5px rgba(0,0,0,0.05);
        margin-bottom: 20px;
    }
    .heading1 {
        font-weight: bold;
        font-size: 1.2em;
        color: #333;
        margin: 15px 0 8px 0;
        border-bottom: 1px solid #eee;
        padding-bottom: 5px;
    }
    .heading2 {
        font-weight: bold;
        font-size: 1.1em;
        color: #444;
        margin: 12px 0 6px 0;
    }
    .heading3 {
        font-weight: bold;
        font-size: 1em;
        color: #555;
        margin: 10px 0 5px 0;
    }
    .score {
        font-weight: bold;
        color: #5cb85c;
    }
    .total-score {
        font-size: 1.1em;
        margin: 5px 0;
    }
    .correct-point {
        color: #5cb85c;
        margin-left: 20px;
        font-size: 0.95em;
        padding: 2px 0;
    }
    .incorrect-point {
        color: #d9534f;
        margin-left: 20px;
        font-size: 0.95em;
        padding: 2px 0;
    }
    .deduction-reason {
        color: #777;
        margin-left: 20px;
        font-size: 0.9em;
        font-style: italic;
        padding: 2px 0;
    }
    .list-item {
        margin-left: 15px;
        font-size: 0.95em;
        padding: 2px 0;
    }
    .math-formula {
        margin-left: 15px;
        font-family: monospace;
        font-size: 0.95em;
    }
    </style>
    """, unsafe_allow_html=True)

    # Format the content with HTML
    formatted_content = content

    # Replace headings
    formatted_content = re.sub(
        r'# (.*?)(\n|$)',
        r'<div class="heading1">\1</div>',
        formatted_content)
    formatted_content = re.sub(
        r'## (.*?)(\n|$)',
        r'<div class="heading2">\1</div>',
        formatted_content)
    formatted_content = re.sub(
        r'### (.*?)(\n|$)',
        r'<div class="heading3">\1</div>',
        formatted_content)

    # Format special headings
    formatted_content = re.sub(
        r'学生答案批改如下:',
        r'<div class="heading1">学生答案批改如下:</div>',
        formatted_content)

    # Format scores
    formatted_content = re.sub(
        r'总分：(\d+)/(\d+)',
        r'<div class="total-score">总分：<span class="score">\1</span>/\2</div>',
        formatted_content)
    formatted_content = re.sub(
        r'(\d+)\s*\([Ii]\)',
        r'<div class="score">\1 分</div>',
        formatted_content)

    # Format step scores
    formatted_content = re.sub(
        r'(\d+)\. 第\d+步：(.*?) - (\d+)/(\d+)',
        r'<div style="font-weight:bold;margin-top:10px;">\1. 第\1步：\2 - <span class="score">\3</span>/\4</div>',
        formatted_content)

    # Format correct/incorrect points
    formatted_content = re.sub(
        r'- ✓ 正确点：(.*?)(\n|$)',
        r'<div class="correct-point">✓ 正确点：\1</div>',
        formatted_content)
    formatted_content = re.sub(
        r'- ✗ 错误点：(.*?)(\n|$)',
        r'<div class="incorrect-point">✗ 错误点：\1</div>',
        formatted_content)
    formatted_content = re.sub(
        r'- 扣分原因：(.*?)(\n|$)',
        r'<div class="deduction-reason">🔍 扣分原因：\1</div>',
        formatted_content)

    # Format subject and question type
    formatted_content = re.sub(
        r'- 科目：(.*?)(\n|$)',
        r'<div style="font-size:0.95em;margin:4px 0;">📚 科目：<span style="color:#333;">\1</span></div>',
        formatted_content)
    formatted_content = re.sub(
        r'- 题目类型：(.*?)(\n|$)',
        r'<div style="font-size:0.95em;margin:4px 0;">📝 题目类型：<span style="color:#333;">\1</span></div>',
        formatted_content)

    # Process remaining list items and general content
    lines = formatted_content.split('\n')
    processed_lines = []

    for line in lines:
        # Skip already processed lines with HTML
        if '<div' in line:
            processed_lines.append(line)
            continue

        # Process list items
        if line.strip().startswith('-') or line.strip().startswith('•'):
            processed_lines.append(f'<div class="list-item">{line}</div>')
        # Process math formulas
        elif line.strip().startswith('∴') or line.strip().startswith('∵'):
            processed_lines.append(f'<div class="math-formula">{line}</div>')
        # Process normal text
        elif line.strip():
            processed_lines.append(
                f'<div style="margin:4px 0;line-height:1.4;">{line}</div>')
        else:
            # Empty line spacing
            processed_lines.append('<div style="height:8px;"></div>')

    # Join processed lines
    formatted_content = ''.join(processed_lines)

    # 基本信息部分与批改结果部分分离处理
    parts = formatted_content.split('<div class="heading1">学生答案批改如下:</div>')
    
    if len(parts) > 1:
        # 显示基本信息部分
        st.markdown(
            f'<div class="result-container">{parts[0]}</div>',
            unsafe_allow_html=True)
        
        # 创建图像与批改结果的交错显示区域
        st.markdown('<div style="margin-top: 20px; border-top: 1px solid #eee; padding-top: 15px;"></div>', unsafe_allow_html=True)
        st.markdown('<div class="heading1">学生答案批改如下:</div>', unsafe_allow_html=True)
        
        # 获取批改内容块
        correction_blocks = parse_correction_blocks(parts[1])
        
        # 对上传的文件按相关性分组
        file_groups = detect_file_groups(files_dict)
        
        last_displayed_image_paths = set() # Track displayed images

        # 处理每个文件组
        for group_idx, group in enumerate(file_groups):
            group_files = {}
            group_title = f"批改组 {group_idx + 1}"
            
            for file_key in group:
                if "student_answer" in file_key:
                    number_match = re.search(r'(\d+)', file_key)
                    if number_match:
                        group_title = f"第 {number_match.group(1)} 题"
                    else:
                        group_title = "学生答案"
                    break
            
            st.subheader(group_title)
            
            current_group_image_paths = set()
            images_to_display_for_current_group = []

            for file_key in group:
                if file_key in files_dict:
                    file_info = files_dict[file_key]
                    file_type_label = "未知类型"
                    type_key = "unknown"
                    if "question" in file_key:
                        file_type_label = "题目"
                        type_key = "question"
                    elif "student_answer" in file_key:
                        file_type_label = "学生答案"
                        type_key = "student_answer"
                    elif "marking_scheme" in file_key:
                        file_type_label = "评分标准"
                        type_key = "marking_scheme"
                    
                    if type_key not in group_files: # Store first encountered of each type
                        group_files[type_key] = {
                            'key': file_key, 'info': file_info, 'type_name': file_type_label
                        }
                    
                    # Collect image paths for current group
                    if file_info and isinstance(file_info, dict) and 'saved_path' in file_info:
                        is_valid, file_path_str = ensure_valid_file_path(file_info['saved_path'])
                        if is_valid:
                            current_group_image_paths.add(file_path_str)
                            images_to_display_for_current_group.append({'path': file_path_str, 'type_name': file_type_label})
            
            # Conditional image display
            if current_group_image_paths and current_group_image_paths != last_displayed_image_paths:
                st.markdown('<div style="display: flex; flex-wrap: wrap; gap: 10px; justify-content: start; margin-bottom: 15px;">', unsafe_allow_html=True)
                # Display unique images for the current group, ordered by type preference
                displayed_in_group = set()
                temp_image_display_list = []
                
                # Prioritize question, then student_answer, then marking_scheme
                for type_to_check in ["question", "student_answer", "marking_scheme"]:
                    for img_info_to_check in images_to_display_for_current_group:
                        if img_info_to_check['type_name'] == group_files.get(type_to_check, {}).get('type_name') and img_info_to_check['path'] not in displayed_in_group:
                            temp_image_display_list.append(img_info_to_check)
                            displayed_in_group.add(img_info_to_check['path'])
                # Add any other images not yet displayed
                for img_info_to_check in images_to_display_for_current_group:
                     if img_info_to_check['path'] not in displayed_in_group:
                        temp_image_display_list.append(img_info_to_check)
                        displayed_in_group.add(img_info_to_check['path'])

                for img_info in temp_image_display_list:
                    st.markdown(f'<div style="flex: 1; min-width: 200px; max-width: 33%; padding: 10px; text-align: center;">', unsafe_allow_html=True)
                    st.markdown(f'<div style="font-weight: bold; margin-bottom: 5px;">{img_info["type_name"]}</div>', unsafe_allow_html=True)
                    try:
                        st.image(img_info['path'], use_column_width=True)
                    except Exception as e:
                        st.info(f"无法显示图片预览: {str(e)}")
                    st.markdown('</div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
                last_displayed_image_paths = current_group_image_paths.copy()
            elif not current_group_image_paths: # Reset if current group has no images
                last_displayed_image_paths = set()
                
            correction_block = None
            if "student_answer" in group_files: # Prioritize matching based on student_answer key
                student_key = group_files["student_answer"]["key"]
                number_match = re.search(r'(\d+)', student_key)
                if number_match:
                    question_num = number_match.group(1)
                    for block in correction_blocks:
                        if f"第{question_num}题" in block or f"第 {question_num} 题" in block:
                            correction_block = block
                            break
            if correction_block is None and len(file_groups) == 1 and len(correction_blocks) == 1:
                correction_block = correction_blocks[0]

            if correction_block:
                st.markdown('<div style="background-color: #f9f9f9; border-radius: 5px; padding: 15px; margin-bottom: 20px;">', unsafe_allow_html=True)
                # Display "批改结果" subheader only if it's not already in the block
                if not re.search(r'<div class="heading2">批改结果</div>', correction_block, re.IGNORECASE) and not re.search(r'批改结果', correction_block[:50], re.IGNORECASE):
                     st.markdown('<div class="heading2">批改结果</div>', unsafe_allow_html=True)
                
                processed_block_lines = []
                for line in correction_block.split('\n'):
                    if line.strip():
                        if re.match(r'\d+\.\s*第\d+步', line):
                            processed_block_lines.append(f'<div style="font-weight: bold; font-size: 1em; margin: 10px 0 5px 0;">{line}</div>')
                        elif "✓ 正确点" in line:
                            processed_block_lines.append(f'<div style="color: #5cb85c; margin-left: 20px; font-size: 0.9em;">{line}</div>')
                        elif "✗ 错误点" in line:
                            processed_block_lines.append(f'<div style="color: #d9534f; margin-left: 20px; font-size: 0.9em;">{line}</div>')
                        elif "扣分原因" in line:
                            processed_block_lines.append(f'<div style="color: #777; margin-left: 20px; font-size: 0.9em; font-style: italic;">{line}</div>')
                        else:
                            processed_block_lines.append(f'<div style="font-size: 0.9em; margin: 3px 0;">{line}</div>')
                st.markdown("".join(processed_block_lines), unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
            
            if group_idx < len(file_groups) - 1:
                st.markdown('<hr style="margin: 30px 0; border-top: 1px dashed #ddd;">', unsafe_allow_html=True)
        
        # 添加总结信息
        st.markdown('<div style="margin-top: 30px; border-top: 1px solid #eee; padding-top: 15px;"></div>', unsafe_allow_html=True)
        st.subheader("总结")
        
        # 提取总分和关键评语
        total_score_match = re.search(r'总分：<span class="score">(\d+)</span>/(\d+)', parts[0])
        if total_score_match:
            score, total = total_score_match.group(1), total_score_match.group(2)
            st.markdown(f"<div style='font-size: 1.3em; margin-bottom: 10px;'>最终得分: <span style='color: #5cb85c; font-weight: bold;'>{score}</span>/{total}</div>", unsafe_allow_html=True)
    else:
        # 如果没有找到分隔标志，直接显示全部内容
        st.markdown(
            f'<div class="result-container">{formatted_content}</div>',
            unsafe_allow_html=True)

    # Download options
    st.subheader("📥 Download Options")
    create_download_options(content, f"record_{record_index}", files_dict)


def ai_correction_page():
    """AI correction management page with enhanced UI"""
    # 检查登录状态
    if not st.session_state.logged_in:
        st.warning("⚠️ 您需要登录才能使用AI批改功能")
        if st.button("点击登录", type="primary"):
            st.session_state.show_login_modal = True
            st.session_state.next_page_after_login = "ai_correction"
            st.session_state.needs_rerun = True
        return

    if 'ai_correction_step' not in st.session_state:
        st.session_state.ai_correction_step = 1

    # 进度条显示当前在哪一步
    steps = ["Step 1: Upload", "Step 2: Settings", "Step 3: Results"]
    current_step = st.session_state.ai_correction_step - 1

    st.markdown(f"""
    <div style="background-color: #f8f9fa; border-radius: 10px; padding: 20px; margin-bottom: 30px;">
        <div style="display: flex; justify-content: space-between; margin-bottom: 15px;">
            <div style="font-weight: bold; color: {"#5cb85c" if current_step >= 0 else "#4a4a4a"}; font-size: 0.9em;">{steps[0]} {" ✓" if current_step > 0 else ""}</div>
            <div style="font-weight: bold; color: {"#5cb85c" if current_step >= 1 else "#4a4a4a"}; font-size: 0.9em;">{steps[1]} {" ✓" if current_step > 1 else ""}</div>
            <div style="font-weight: bold; color: {"#5cb85c" if current_step >= 2 else "#4a4a4a"}; font-size: 0.9em;">{steps[2]}</div>
        </div>
        <div style="height: 6px; background-color: #e9ecef; border-radius: 3px; position: relative;">
            <div style="height: 100%; width: {(current_step + 1) * 33.33}%; background-color: #5cb85c; border-radius: 3px;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.ai_correction_step == 1:
        st.header("Step 1: Upload Files")

        # 添加文件格式提示
        st.markdown("""
        <style>
        .tooltip {
          position: relative;
          display: inline-block;
        }
        .tooltip .tooltiptext {
          visibility: hidden;
          width: 320px;
          background-color: #555;
          color: #fff;
          text-align: left;
          border-radius: 6px;
          padding: 10px;
          position: absolute;
          z-index: 1;
          bottom: 125%; /* 上方显示 */
          left: 50%;
          margin-left: -160px;
          opacity: 0;
          transition: opacity 0.3s;
          font-size: 0.95em;
        }
        .tooltip:hover .tooltiptext {
          visibility: visible;
          opacity: 1;
        }
        </style>
        <div class="tooltip">
          <button style="padding: 6px 18px; border-radius: 5px; border: none; background: #4b8bf4; color: #fff; font-weight: bold; cursor: pointer;">
            File Requirements
          </button>
          <span class="tooltiptext">
            <b>Supported formats:</b> PDF, JPG, JPEG, PNG, JSON<br>
            <b>Maximum file size:</b> 5MB per file<br>
            <b>Images should be clear and readable</b><br>
            <b>Student answer file is required</b>
          </span>
        </div>
        """, unsafe_allow_html=True)

        question_files = st.file_uploader(
            "Upload question(s) (optional)",
            type=[
                "pdf",
                "jpg",
                "jpeg",
                "png"],
            key="question_file",
            accept_multiple_files=True)
        student_answer_files = st.file_uploader(
            "Upload student answer(s)",
            type=[
                "pdf",
                "jpg",
                "jpeg",
                "png"],
            key="student_answer_file",
            accept_multiple_files=True)
        marking_scheme_files = st.file_uploader(
            "Upload marking scheme(s) (optional)",
            type=[
                "pdf",
                "jpg",
                "jpeg",
                "png",
                "json"],
            key="marking_scheme_file",
            accept_multiple_files=True)

        if st.button("Next Step"):
            if not student_answer_files or len(student_answer_files) == 0:
                st.warning("Please upload at least one student answer file.")
            else:
                st.session_state.question_files = question_files
                st.session_state.student_answer_files = student_answer_files
                st.session_state.marking_scheme_files = marking_scheme_files
                st.session_state.ai_correction_step = 2
                st.session_state.needs_rerun = True

    elif st.session_state.ai_correction_step == 2:
        # 获取会话中存储的文件
        question_files = st.session_state.get("question_files", [])
        student_answer_files = st.session_state.get("student_answer_files", [])
        marking_scheme_files = st.session_state.get("marking_scheme_files", [])

        # 检查是否有必需的文件
        if not student_answer_files or len(student_answer_files) == 0:
            st.error("No student answer files found. Please go back and upload.")
            if st.button("Back to Upload"):
                st.session_state.ai_correction_step = 1
                st.session_state.needs_rerun = True
            return

        st.header("Step 2: Process")

        # 添加批改设置
        st.markdown("### Correction Settings")

        # 简化的严格程度选择
        strictness_level = st.radio(
            "Grading Strictness",
            ["宽松", "中等", "严格"],
            index=1,  # 默认选择中等
            horizontal=True
        )

        # 简化的语言选择
        language = st.radio(
            "Output Language",
            ["中文", "English"],
            index=0,  # 默认选择中文
            horizontal=True,
            format_func=lambda x: x
        )

        # 将radio的英文值转换为API所需的代码
        if language == "English":
            language = "en"
        else:
            language = "zh"

        # 提交按钮区域，增加醒目程度
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("Back", use_container_width=True):
                st.session_state.ai_correction_step = 1
                st.session_state.needs_rerun = True
        with col2:
            process_button = st.button(
                "📝 Process Correction",
                type="primary",
                use_container_width=True)

        # 检查是否点击了处理按钮
        if process_button:
            # 保存设置到会话状态
            st.session_state.strictness_level = strictness_level
            st.session_state.language = language
            st.session_state.ai_correction_step = 3
            st.session_state.needs_rerun = True

    elif st.session_state.ai_correction_step == 3:
        st.header("Step 3: Results")

        # 获取所有必要的数据
        question_files = st.session_state.get("question_files", [])
        student_answer_files = st.session_state.get("student_answer_files", [])
        marking_scheme_files = st.session_state.get("marking_scheme_files", [])
        strictness_level = st.session_state.get("strictness_level", "中等")
        language = st.session_state.get("language", "zh")

        # 创建用户目录
        user_dir = UPLOAD_DIR / st.session_state.current_user
        user_dir.mkdir(exist_ok=True)

        # 读取用户数据
        user_data = read_user_data()

        # 始终显示已上传文件的预览，无论生成结果的状态如何
        if student_answer_files or question_files or marking_scheme_files:
            st.subheader("Uploaded Files Preview")
            cols = st.columns(3)

            with cols[0]:
                if question_files:
                    st.write("Question Files:")
                    for q in question_files:
                        try:
                            st.image(q, caption=q.name, width=200)
                        except BaseException:
                            st.info(f"Preview not available: {q.name}")

            with cols[1]:
                if student_answer_files:
                    st.write("Student Answer Files:")
                    for sa in student_answer_files:
                        try:
                            st.image(sa, caption=sa.name, width=200)
                        except BaseException:
                            st.info(f"Preview not available: {sa.name}")

            with cols[2]:
                if marking_scheme_files:
                    st.write("Marking Scheme Files:")
                    for ms in marking_scheme_files:
                        try:
                            st.image(ms, caption=ms.name, width=200)
                        except BaseException:
                            st.info(f"Preview not available: {ms.name}")

        # 显示处理中的消息和动画
        if 'correction_success' not in st.session_state:
            with st.spinner("AI analyzing student answer. This may take a moment..."):
                # 保存上传的文件（如果存在）
                user_files = {}

                if question_files:
                    for i, question in enumerate(question_files):
                        question_path = save_uploaded_file(
                            user_dir, question, f"question_{i + 1}", user_data)
                        user_files[f"question_{i + 1}"] = {
                            "filename": question.name, "saved_path": str(question_path)}

                if student_answer_files:
                    for i, student_answer in enumerate(student_answer_files):
                        student_answer_path = save_uploaded_file(
                            user_dir, student_answer, f"student_answer_{i + 1}", user_data)
                        user_files[f"student_answer_{i + 1}"] = {
                            "filename": student_answer.name, "saved_path": str(student_answer_path)}

                if marking_scheme_files:
                    for i, marking_scheme in enumerate(marking_scheme_files):
                        marking_scheme_path = save_uploaded_file(
                            user_dir, marking_scheme, f"marking_scheme_{i + 1}", user_data)
                        user_files[f"marking_scheme_{i + 1}"] = {
                            "filename": marking_scheme.name, "saved_path": str(marking_scheme_path)}

                try:
                    # 创建图像文件列表
                    image_files = []
                    if question_files:
                        image_files.extend(question_files)
                        for q in question_files:
                            q.seek(0)  # 重置文件指针

                    if student_answer_files:
                        image_files.extend(student_answer_files)
                        for sa in student_answer_files:
                            sa.seek(0)  # 重置文件指针

                    if marking_scheme_files:
                        image_files.extend(marking_scheme_files)
                        for ms in marking_scheme_files:
                            ms.seek(0)  # 重置文件指针

                    logging.info(
                        f"Preparing to call API with {
                            len(image_files)} files and strictness level: {strictness_level}, language: {language}")

                    # 确保语言参数有效
                    if language not in ["zh", "en"]:
                        language = "zh"
                        logging.warning(
                            f"Invalid language parameter: {language}, defaulting to 'zh'")

                    # 调用API进行批改
                    api_result = correction_with_image_marking_scheme(
                        *image_files,
                        strictness_level=strictness_level,
                        language=language
                    )

                    if api_result and isinstance(api_result, str):
                        # 保存结果到会话状态
                        st.session_state.correction_success = True
                        st.session_state.correction_result = api_result

                        # 保存 user_files 到会话状态
                        st.session_state.user_files = user_files

                        # 保存结果到用户记录
                        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        correction_record = {
                            "timestamp": timestamp,
                            "content": api_result,
                            "files": user_files,
                            "settings": {
                                "strictness_level": strictness_level,
                                "language": language
                            }
                        }

                        user_data[st.session_state.current_user]["records"].append(
                            correction_record)
                        save_user_data(user_data)
                    else:
                        st.session_state.correction_success = False
                        st.session_state.correction_error = "API returned an invalid result"
                except Exception as e:
                    st.session_state.correction_success = False
                    st.session_state.correction_error = str(e)
                    logging.error(f"Correction processing error: {str(e)}")

                # 标记需要刷新页面以显示结果
                st.session_state.needs_rerun = True

        # 显示结果
        if hasattr(
                st.session_state,
                'correction_success') and st.session_state.correction_success:
            st.success("✅ Correction processed successfully!")

            st.markdown(
                '<div style="background-color: #f8f9fa; border-radius: 10px; padding: 20px; margin-bottom: 30px;">',
                unsafe_allow_html=True)
            st.markdown(
                '<div style="font-size: 1.1em; font-weight: bold; color: #333; margin-bottom: 15px;">Correction Results</div>',
                unsafe_allow_html=True)

            # 增强结果显示区域
            st.markdown(
                '<div style="background-color: white; border-radius: 5px; padding: 20px; border-left: 5px solid #5cb85c; box-shadow: 0 2px 5px rgba(0,0,0,0.1);">',
                unsafe_allow_html=True)

            # 优化结果的显示方式，替换纯文本为格式化内容
            result_text = st.session_state.correction_result

            # 首先处理主要的标题格式
            result_text = re.sub(
                r'# (.*?)(\n|$)',
                r'<div style="font-weight: bold; font-size: 1.1em; color: #333; margin: 10px 0 5px 0; border-bottom: 1px solid #eee; padding-bottom: 5px;">\1</div>',
                result_text)
            result_text = re.sub(
                r'## (.*?)(\n|$)',
                r'<div style="font-weight: bold; font-size: 1em; color: #444; margin: 8px 0 4px 0;">\1</div>',
                result_text)
            result_text = re.sub(
                r'### (.*?)(\n|$)',
                r'<div style="font-weight: bold; font-size: 0.95em; color: #555; margin: 6px 0 3px 0;">\1</div>',
                result_text)

            # 处理其他特定标题格式
            result_text = re.sub(
                r'学生答案批改如下:',
                r'<div style="font-weight: bold; font-size: 1em; color: #333;">学生答案批改如下:</div>',
                result_text)

            # 处理分数显示
            result_text = re.sub(
                r'总分：(\d+)/(\d+)',
                r'<div style="font-size: 0.95em; margin: 5px 0;">总分：<span style="font-weight: bold; color: #5cb85c;">\1</span>/\2</div>',
                result_text)
            result_text = re.sub(
                r'(\d+)\s*\([Ii]\)',
                r'<div style="font-size: 1em; font-weight: bold; color: #5cb85c; margin: 8px 0;">\1 分</div>',
                result_text)

            # 处理步骤评分
            result_text = re.sub(
                r'(\d+)\. 第\d+步：(.*?) - (\d+)/(\d+)',
                r'<div style="font-size: 0.95em; margin: 10px 0 5px 0;"><span style="font-weight: bold;">\1. 第\1步：\2</span> - <span style="color: #5cb85c; font-weight: bold;">\3</span>/\4</div>',
                result_text)

            # 处理正确和错误点
            result_text = re.sub(
                r'- ✓ 正确点：(.*?)(\n|$)',
                r'<div style="color: #5cb85c; margin-left: 20px; font-size: 0.9em;">✓ 正确点：\1</div>',
                result_text)
            result_text = re.sub(
                r'- ✗ 错误点：(.*?)(\n|$)',
                r'<div style="color: #d9534f; margin-left: 20px; font-size: 0.9em;">✗ 错误点：\1</div>',
                result_text)
            result_text = re.sub(
                r'- 扣分原因：(.*?)(\n|$)',
                r'<div style="color: #777; margin-left: 20px; font-size: 0.9em;">🔍 扣分原因：\1</div>',
                result_text)

            # 科目和题型信息
            result_text = re.sub(
                r'- 科目：(.*?)(\n|$)',
                r'<div style="font-size: 0.9em; margin: 3px 0;">📚 科目：<span style="color: #333;">\1</span></div>',
                result_text)
            result_text = re.sub(
                r'- 题目类型：(.*?)(\n|$)',
                r'<div style="font-size: 0.9em; margin: 3px 0;">📝 题目类型：<span style="color: #333;">\1</span></div>',
                result_text)

            # 处理一般列表项
            lines = result_text.split('\n')
            formatted_lines = []
            for line in lines:
                # 跳过已经处理过的行
                if '<div' in line:
                    formatted_lines.append(line)
                    continue

                # 处理列表项
                if line.strip().startswith('-') or line.strip().startswith('•'):
                    line = f'<div style="margin-left: 15px; font-size: 0.9em;">{line}</div>'
                # 处理缩进的数学公式
                elif line.strip().startswith('∴') or line.strip().startswith('∵'):
                    line = f'<div style="margin-left: 15px; font-family: monospace; font-size: 0.9em;">{line}</div>'
                # 处理普通文本
                elif line.strip():
                    line = f'<div style="font-size: 0.9em; line-height: 1.4; margin: 3px 0;">{line}</div>'
                else:
                    line = '<div style="height: 8px;"></div>'  # 空行
                formatted_lines.append(line)

            result_text = ''.join(formatted_lines)

            # 应用一致的基础字体大小和行高
            result_text = f'<div style="font-size: 14px; line-height: 1.4;">{result_text}</div>'

            # 基本信息部分先显示
            parts = result_text.split('<div style="font-weight: bold; font-size: 1em; color: #333;">学生答案批改如下:</div>')
            
            if len(parts) > 1:
                # 显示基本信息部分
                st.markdown(parts[0], unsafe_allow_html=True)
                
                # 获取用户文件进行显示
                user_files_from_session = st.session_state.get('user_files', {}) # Corrected Indentation
                
                # 创建图像与批改结果的交错显示区域
                st.markdown('<div style="margin-top: 20px; border-top: 1px solid #eee; padding-top: 15px;"></div>', unsafe_allow_html=True)
                st.markdown('<div style="font-weight: bold; font-size: 1em; color: #333;">学生答案批改如下:</div>', unsafe_allow_html=True)

                correction_blocks = parse_correction_blocks(parts[1])
                file_groups = detect_file_groups(user_files_from_session)
                
                last_displayed_image_paths = set() # Track displayed images

                for group_idx, group in enumerate(file_groups):
                    group_files = {}
                    group_title = f"批改组 {group_idx + 1}"
                    
                    # Corrected logic for setting group_title
                    for file_key in group:
                        if "student_answer" in file_key:
                            number_match = re.search(r'(\d+)', file_key)
                            if number_match:
                                group_title = f"第 {number_match.group(1)} 题"
                            else:
                                group_title = "学生答案"
                            break 
                    
                    st.subheader(group_title)
                    
                    current_group_image_paths = set()
                    images_to_display_for_current_group = []

                    for file_key in group:
                        if file_key in user_files_from_session:
                            file_info = user_files_from_session[file_key]
                            file_type_label = "未知类型"
                            # Determine file type for display label and internal key
                            type_key = "unknown"
                            if "question" in file_key:
                                file_type_label = "题目"
                                type_key = "question"
                            elif "student_answer" in file_key:
                                file_type_label = "学生答案"
                                type_key = "student_answer"
                            elif "marking_scheme" in file_key:
                                file_type_label = "评分标准"
                                type_key = "marking_scheme"
                            
                            # Store first encountered file of each type for the group
                            if type_key not in group_files:
                                group_files[type_key] = {
                                    'key': file_key, 'info': file_info, 'type_name': file_type_label
                                }

                            # Collect image paths for current group to check for redundancy
                            if file_info and isinstance(file_info, dict) and 'saved_path' in file_info:
                                is_valid, file_path_str = ensure_valid_file_path(file_info['saved_path'])
                                if is_valid:
                                    current_group_image_paths.add(file_path_str)
                                    # Add to a list that preserves order and includes type for display
                                    images_to_display_for_current_group.append({'path': file_path_str, 'type_name': file_type_label})
                    
                    # Conditional image display
                    if current_group_image_paths and current_group_image_paths != last_displayed_image_paths:
                        st.markdown('<div style="display: flex; flex-wrap: wrap; gap: 10px; justify-content: start; margin-bottom: 15px;">', unsafe_allow_html=True)
                        
                        # Display unique images for the current group, ordered by type preference
                        displayed_in_group = set()
                        temp_image_display_list = []
                        
                        # Prioritize question, then student_answer, then marking_scheme
                        for type_to_check in ["question", "student_answer", "marking_scheme"]:
                            for img_info_to_check in images_to_display_for_current_group:
                                # Check if this image type matches the current priority and hasn't been added from this group
                                if img_info_to_check['type_name'] == group_files.get(type_to_check, {}).get('type_name') and img_info_to_check['path'] not in displayed_in_group:
                                    temp_image_display_list.append(img_info_to_check)
                                    displayed_in_group.add(img_info_to_check['path'])
                        
                        # Add any other images from the group not yet displayed (e.g., multiple 'unknown' or if type not in group_files)
                        for img_info_to_check in images_to_display_for_current_group:
                             if img_info_to_check['path'] not in displayed_in_group:
                                temp_image_display_list.append(img_info_to_check)
                                displayed_in_group.add(img_info_to_check['path'])

                        for img_info in temp_image_display_list:
                            st.markdown(f'<div style="flex: 1; min-width: 200px; max-width: 33%; padding: 10px; text-align: center;">', unsafe_allow_html=True)
                            st.markdown(f'<div style="font-weight: bold; margin-bottom: 5px;">{img_info["type_name"]}</div>', unsafe_allow_html=True)
                            try:
                                st.image(img_info['path'], use_column_width=True)
                            except Exception as e:
                                st.info(f"无法显示图片预览: {str(e)}")
                            st.markdown('</div>', unsafe_allow_html=True)
                        st.markdown('</div>', unsafe_allow_html=True)
                        last_displayed_image_paths = current_group_image_paths.copy() # Update for next iteration
                    elif not current_group_image_paths: # Reset if current group has no images, so next group with images will show them
                        last_displayed_image_paths = set()
                        
                    correction_block = None
                    # Try to match correction block more robustly
                    if "student_answer" in group_files: # Prioritize matching based on student_answer key
                        student_key = group_files["student_answer"]["key"]
                        number_match = re.search(r'(\d+)', student_key)
                        if number_match:
                            question_num = number_match.group(1)
                            for block in correction_blocks:
                                if f"第{question_num}题" in block or f"第 {question_num} 题" in block:
                                    correction_block = block
                                    break
                        # If no specific number match, but it's the only group and only block
                        if correction_block is None and len(file_groups) == 1 and len(correction_blocks) == 1:
                            correction_block = correction_blocks[0]
                    elif correction_blocks: # Fallback if no student_answer or if only one block for multiple groups
                         if group_idx < len(correction_blocks): # Try to match by index if multiple blocks
                            correction_block = correction_blocks[group_idx]
                         elif len(correction_blocks) == 1: # If only one block, use it for all groups
                            correction_block = correction_blocks[0]

                    if correction_block:
                        st.markdown('<div style="background-color: #f9f9f9; border-radius: 5px; padding: 15px; margin-bottom: 20px;">', unsafe_allow_html=True)
                        # Display "批改结果" subheader only if it's not already in the block (common in parsed blocks)
                        if not re.search(r'<div style="font-weight: bold; font-size: 1em; color: #444; margin: 8px 0 4px 0;">批改结果</div>', correction_block, re.IGNORECASE) and not re.search(r'批改结果', correction_block[:50], re.IGNORECASE): # Check common variations
                             st.markdown('<div style="font-weight: bold; font-size: 1em; color: #444; margin: 8px 0 4px 0;">批改结果</div>', unsafe_allow_html=True)
                        
                        # Process and display correction lines (same as before)
                        processed_block_lines = []
                        for line in correction_block.split('\n'): # Assuming blocks are newline separated
                            if line.strip(): # Ensure line is not just whitespace
                                # Apply styling based on content
                                if re.match(r'\d+\.\s*第\d+步', line): # Step scoring
                                    processed_block_lines.append(f'<div style="font-weight: bold; font-size: 1em; margin: 10px 0 5px 0;">{line}</div>')
                                elif "✓ 正确点" in line:
                                    processed_block_lines.append(f'<div style="color: #5cb85c; margin-left: 20px; font-size: 0.9em;">{line}</div>')
                                elif "✗ 错误点" in line:
                                    processed_block_lines.append(f'<div style="color: #d9534f; margin-left: 20px; font-size: 0.9em;">{line}</div>')
                                elif "扣分原因" in line:
                                    processed_block_lines.append(f'<div style="color: #777; margin-left: 20px; font-size: 0.9em; font-style: italic;">{line}</div>')
                                else: # Default line styling
                                    processed_block_lines.append(f'<div style="font-size: 0.9em; margin: 3px 0;">{line}</div>')
                        st.markdown("".join(processed_block_lines), unsafe_allow_html=True)
                        st.markdown('</div>', unsafe_allow_html=True)
                    
                    if group_idx < len(file_groups) - 1:
                        st.markdown('<hr style="margin: 30px 0; border-top: 1px dashed #ddd;">', unsafe_allow_html=True)
            
            else: # Corrected Indentation for the outer else
                st.markdown(result_text, unsafe_allow_html=True)

        elif hasattr(st.session_state, 'correction_success') and not st.session_state.correction_success:
            error_message = getattr(
                st.session_state,
                'correction_error',
                "Unknown error occurred")
            st.error(f"❌ Failed to process correction: {error_message}")

            # 提供更多具体的错误指导
            if "timeout" in error_message.lower() or "connection" in error_message.lower():
                st.info(
                    "Failed to connect to API server. Please check your network connection and try again later.")
            elif "api key" in error_message.lower():
                st.info(
                    "API key issue. Please contact administrator to check API settings.")
            elif "'NoneType' object has no attribute" in error_message:
                st.info(
                    "There was an issue processing one of the uploaded files. Please ensure your files are valid images or PDFs.")
            elif "Failed to read uploaded file" in error_message:
                st.info(
                    "Unable to read one of the uploaded files. Please ensure your files are not corrupted.")

        # 添加操作按钮
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Start New Correction", use_container_width=True):
                # 清除结果和错误
                if 'correction_success' in st.session_state:
                    del st.session_state.correction_success
                if 'correction_result' in st.session_state:
                    del st.session_state.correction_result
                if 'correction_error' in st.session_state:
                    del st.session_state.correction_error
                # 返回第一步
                st.session_state.ai_correction_step = 1
                st.session_state.needs_rerun = True

        with col2:
            if st.button("⚙️ Adjust Settings", use_container_width=True):
                # 保留上传的文件但回到设置页面
                st.session_state.ai_correction_step = 2
                # 清除结果但不清除上传的文件
                if 'correction_success' in st.session_state:
                    del st.session_state.correction_success
                if 'correction_result' in st.session_state:
                    del st.session_state.correction_result
                if 'correction_error' in st.session_state:
                    del st.session_state.correction_error
                st.session_state.needs_rerun = True

    # 添加页脚
    st.markdown("""
    <div style="text-align: center; margin-top: 30px; padding-top: 15px; border-top: 1px solid #eee; color: #666; font-size: 0.8em;">
        AI Guru Correction System • Powered by advanced AI • © 2025
    </div>
    """, unsafe_allow_html=True)

# 新增辅助函数用于保存上传的文件


def save_uploaded_file(user_dir, uploaded_file, file_type, user_data):
    """
    保存上传的文件并更新用户记录

    参数:
    user_dir: Path对象，用户目录路径
    uploaded_file: UploadedFile对象，上传的文件
    file_type: str，文件类型
    user_data: dict，用户数据字典

    返回:
    Path对象，保存的文件路径
    """
    file_path = user_dir / uploaded_file.name
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # 更新用户记录 你
    file_size = uploaded_file.size / 1024
    record = {
        "filename": uploaded_file.name,
        "upload_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "file_size": round(file_size, 2),
        "file_type": file_type,
        "processing_result": "Uploaded"
    }

    if st.session_state.current_user in user_data:
        user_data[st.session_state.current_user]["records"].append(record)
        save_user_data(user_data)

    return file_path

# 添加密码哈希函数


def hash_password(password):
    """对密码进行安全哈希处理"""
    return hashlib.sha256(password.encode()).hexdigest()

# 修改主函数添加注册功能


def main():
    # 初始化会话状态 - 确保在应用启动时正确初始化所有需要的状态变量
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'current_user' not in st.session_state:
        st.session_state.current_user = None
    if 'page' not in st.session_state:
        st.session_state.page = 'main_menu'
    if 'show_register' not in st.session_state:
        st.session_state.show_register = False
    if 'ai_correction_step' not in st.session_state:
        st.session_state.ai_correction_step = 1
    if 'previous_page' not in st.session_state:
        st.session_state.previous_page = 'main_menu'
    if 'show_login_modal' not in st.session_state:
        st.session_state.show_login_modal = False
    if 'next_page_after_login' not in st.session_state:
        st.session_state.next_page_after_login = None
    if 'needs_rerun' not in st.session_state:
        st.session_state.needs_rerun = False
    if 'history_page' not in st.session_state:
        st.session_state.history_page = 0
    if 'selected_record' not in st.session_state:
        st.session_state.selected_record = None
    if 'selected_record_index' not in st.session_state:
        st.session_state.selected_record_index = None

    # 打印当前页面状态，用于调试
    logging.info(f"Current page: {st.session_state.page}")
    if st.session_state.page == "ai_correction":
        logging.info(
            f"Current AI correction step: {
                st.session_state.ai_correction_step}")

    # 侧边栏导航（始终显示，但功能根据登录状态调整）
    with st.sidebar:
        st.title("🎓 AI Guru")

        # 显示用户信息或登录按钮
        if st.session_state.logged_in:
            st.write(f"Welcome, {st.session_state.current_user}!")
            if st.button("🚪 Logout"):
                st.session_state.logged_in = False
                st.session_state.current_user = None
                # 如果用户在需要登录的页面，则返回主菜单
                if st.session_state.page in [
                    "ai_correction",
                    "file_management",
                    "history_records",
                        "record_detail"]:
                    st.session_state.page = "main_menu"
                st.session_state.needs_rerun = True
        else:
            st.info("You are not logged in")
            if st.button("🔑 Login / Register"):
                st.session_state.show_login_modal = True
                st.session_state.needs_rerun = True

        # 导航菜单
        st.subheader("📍 Navigation")
        menu_options = {
            "main_menu": "🏠 Main Menu",
            "account_management": "👤 Account Management",
            "file_management": "📁 File Management",
            "ai_correction": "🤖 AI Correction",
        }

        # 不在主要菜单中显示历史记录和记录详情页面，但保持它们的导航功能
        hidden_pages = ["history_records", "record_detail"]

        # 根据当前页面设置默认选择，如果是隐藏页面则显示它们的父页面
        current_page_key = st.session_state.page
        if current_page_key in hidden_pages:
            if current_page_key == "history_records" or current_page_key == "record_detail":
                display_page_key = "file_management"
            else:
                display_page_key = "main_menu"
        else:
            display_page_key = current_page_key

        current_page_value = menu_options.get(
            display_page_key, menu_options["main_menu"])

        # 记录导航选择之前的页面
        st.session_state.previous_page = current_page_key

        # 修复在回调中调用st.rerun()的问题
        # 首先定义回调函数以更新会话状态而不是直接重新运行
        def on_page_change():
            selected_option = st.session_state.page_selector
            page_key = list(menu_options.keys())[list(
                menu_options.values()).index(selected_option)]

            # 只有在页面真正变化时才更新
            if page_key != st.session_state.previous_page:
                # 检查AI批改页面的访问权限
                if page_key == "ai_correction" and not st.session_state.logged_in:
                    st.session_state.show_login_modal = True
                    # 不改变页面，等待登录
                else:
                    st.session_state.page = page_key
                    # 如果切换到AI批改页面，重置步骤
                    if page_key == "ai_correction":
                        st.session_state.ai_correction_step = 1

                # 添加一个标志，在主循环中检测并执行重新运行
                st.session_state.needs_rerun = True

        # 使用key参数和on_change回调来跟踪选择变化
        selected_page = st.radio(
            "Go to:",
            list(menu_options.values()),
            index=list(menu_options.values()).index(current_page_value),
            key="page_selector",
            on_change=on_page_change
        )

    # 登录模态窗口 - 当show_login_modal为True时显示
    if st.session_state.show_login_modal:
        # 使用container创建"模态"效果
        modal_container = st.container()

        with modal_container:
            # 模态窗口的背景和样式
            st.markdown("""
            <style>
            .login-modal {
                background-color: white;
                padding: 20px;
                border-radius: 10px;
                box-shadow: 0 4px 12px rgba(0,0,0,0.15);
                margin: 10px;
                border: 1px solid #ddd;
            }
            </style>
            """, unsafe_allow_html=True)

            # 模态窗口内容
            with st.container():
                st.markdown(
                    '<div class="login-modal">',
                    unsafe_allow_html=True)

                st.subheader("🔐 User Authentication")

                # 关闭按钮
                if st.button("✖️ Close", key="close_modal"):
                    st.session_state.show_login_modal = False
                    st.session_state.needs_rerun = True

                # 切换登录/注册按钮
                col1, col2 = st.columns(2)
                with col1:
                    if st.button(
                        "Login",
                        use_container_width=True,
                            type="primary" if not st.session_state.show_register else "secondary"):
                        st.session_state.show_register = False
                        st.session_state.needs_rerun = True
                with col2:
                    if st.button(
                        "Register",
                        use_container_width=True,
                            type="primary" if st.session_state.show_register else "secondary"):
                        st.session_state.show_register = True
                        st.session_state.needs_rerun = True

                # 根据状态显示登录或注册表单
                if st.session_state.show_register:
                    # 注册表单
                    with st.form("register_form"):
                        st.subheader("📝 Create New Account")
                        new_username = st.text_input("Username")
                        new_password = st.text_input(
                            "Password", type="password")
                        confirm_password = st.text_input(
                            "Confirm Password", type="password")
                        email = st.text_input("Email (optional)")

                        register_submitted = st.form_submit_button("Register")

                        if register_submitted:
                            # 进行表单验证
                            if not new_username or not new_password:
                                st.error("Username and password are required.")
                            elif new_password != confirm_password:
                                st.error("Passwords do not match.")
                            else:
                                # 检查用户名是否已存在
                                user_data = read_user_data()
                                if new_username in user_data:
                                    st.error(
                                        "Username already exists. Please choose another one.")
                                else:
                                    # 创建新用户
                                    user_data[new_username] = {
                                        "password": hash_password(new_password),
                                        "email": email,
                                        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                                        "records": []}
                                    save_user_data(user_data)

                                    # 提示成功并自动设置为登录状态
                                    st.success(
                                        "Registration successful! You can now log in.")
                                    st.session_state.show_register = False
                                    st.session_state.needs_rerun = True
                else:
                    # 登录表单
                    with st.form("login_form"):
                        st.subheader("👤 Login to Your Account")
                        username = st.text_input("Username")
                        password = st.text_input("Password", type="password")
                        login_submitted = st.form_submit_button("Login")

                        if login_submitted:
                            if not username or not password:
                                st.error(
                                    "Please enter both username and password.")
                            else:
                                # 检查测试账户
                                if username in TEST_ACCOUNTS and TEST_ACCOUNTS[
                                        username]['password'] == password:
                                    st.session_state.logged_in = True
                                    st.session_state.current_user = username
                                    st.session_state.show_login_modal = False

                                    # 登录后导航到指定页面（如果有）
                                    if 'next_page_after_login' in st.session_state and st.session_state.next_page_after_login:
                                        st.session_state.page = st.session_state.next_page_after_login
                                        # 清除导航目标，避免重复使用
                                        st.session_state.next_page_after_login = None

                                    st.success("Login successful!")
                                    st.session_state.needs_rerun = True
                                else:
                                    # 检查注册用户
                                    user_data = read_user_data()
                                    if username in user_data and user_data[username].get(
                                            'password') == hash_password(password):
                                        st.session_state.logged_in = True
                                        st.session_state.current_user = username
                                        st.session_state.show_login_modal = False

                                        # 登录后导航到指定页面（如果有）
                                        if 'next_page_after_login' in st.session_state and st.session_state.next_page_after_login:
                                            st.session_state.page = st.session_state.next_page_after_login
                                            # 清除导航目标，避免重复使用
                                            st.session_state.next_page_after_login = None

                                        st.success("Login successful!")
                                        st.session_state.needs_rerun = True
                                    else:
                                        st.error(
                                            "Invalid username or password.")

                # 添加一个忘记密码的链接
                st.markdown("---")
                st.markdown(
                    "<div style='text-align: center'>Forgot your password? Contact administrator.</div>",
                    unsafe_allow_html=True)
                # 关闭login-modal div
                st.markdown('</div>', unsafe_allow_html=True)

    # 页面路由
    if st.session_state.page == "file_management":
        file_management_page()
    elif st.session_state.page == "ai_correction":
        # 检查登录状态，如果未登录则提示登录
        if not st.session_state.logged_in:
            st.warning("You need to log in to use the AI Correction feature.")
            if st.button("Log in now"):
                st.session_state.show_login_modal = True
                st.session_state.needs_rerun = True
        else:
            ai_correction_page()
    elif st.session_state.page == "account_management":
        account_management_page()
    elif st.session_state.page == "history_records":
        history_records_page()
    elif st.session_state.page == "record_detail":
        record_detail_page()
    else:  # main menu
        main_page()

    # 检查是否需要重新运行页面 - 这需要在所有UI组件渲染后执行
    if st.session_state.needs_rerun:
        st.session_state.needs_rerun = False  # 重置标志，避免循环
        st.rerun()


def main_page():
    """Main landing page - default entry point for the application"""
    # 主页显示网站介绍和功能导航
    st.markdown("""
    <div style="text-align: center; margin-bottom: 30px;">
        <h1>
            <span style="font-size: 48px; color: #FF4B4B;">AI</span>
            <span style="font-size: 42px; color: #444;">Guru</span>
        </h1>
        <p style="font-size: 18px; color: #666; margin-top: -10px;">智能批改系统 | 提升评分效率和准确性</p>
    </div>
    """, unsafe_allow_html=True)

    # 如果用户未登录，显示欢迎消息和登录提示
    if not st.session_state.logged_in:
        st.info(
            "👋 Welcome to AI Guru! Sign in to use all features or explore the interface as a guest.")

    # 网站介绍
    st.markdown("""
    <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; margin-bottom: 25px;">
        <h3 style="color: #444; margin-bottom: 15px;">🔍 关于 AI Guru</h3>
        <p style="color: #555;">
            AI Guru 是一款针对教育领域的智能批改助手，利用先进的人工智能技术，帮助教师和教育工作者快速、准确地批改学生作业。
            系统支持中英双语，能够根据评分标准对学生答案进行全面分析，提供详细的评分报告和改进建议。
        </p>
    </div>
    """, unsafe_allow_html=True)

    # 强调AI批改功能
    st.markdown("""
    <div style="background-color: #e8f4ff; padding: 25px; border-radius: 10px; margin-bottom: 25px; border-left: 5px solid #0066cc;">
        <h2 style="color: #0066cc; margin-bottom: 15px;">🤖 AI 智能批改</h2>
        <p style="color: #444; font-size: 16px; margin-bottom: 15px;">
            通过上传题目、学生答案和评分标准，获取详细的AI批改结果。系统能够识别关键点、分析错误，并提供专业的评分建议。
        </p>
    """, unsafe_allow_html=True)

    # 批改功能按钮 - 根据登录状态调整行为
    if st.session_state.logged_in:
        st.markdown("""
        <div style="text-align: center; margin: 20px 0;">
        """, unsafe_allow_html=True)

        if st.button("🚀 开始批改", type="primary", use_container_width=True):
            logging.info("从主页点击了开始批改按钮 - 已登录用户")
            st.session_state.page = "ai_correction"
            st.session_state.ai_correction_step = 1
            st.session_state.needs_rerun = True

        st.markdown("</div>", unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="text-align: center; margin: 20px 0;">
        """, unsafe_allow_html=True)

        if st.button("🔑 登录并开始批改", type="primary", use_container_width=True):
            logging.info("从主页点击了登录并开始批改按钮 - 未登录用户")
            st.session_state.show_login_modal = True
            # 登录后跳转到AI批改页面
            st.session_state.next_page_after_login = "ai_correction"
            st.session_state.needs_rerun = True

        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

    # 创建两栏布局
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("""
        <div style="background-color: #f0f7f0; padding: 20px; border-radius: 10px; height: 100%;">
            <h3 style="color: #2e8b57; margin-bottom: 15px;">✨ 主要功能</h3>
            <ul style="color: #444; padding-left: 20px;">
                <li><b>AI 智能评分</b> - 自动分析并评分学生答案</li>
                <li><b>双语支持</b> - 中英文智能评分与反馈</li>
                <li><b>多文件处理</b> - 批量处理多个学生答案</li>
                <li><b>灵活评分</b> - 可调整评分严格程度</li>
                <li><b>详细分析</b> - 提供全面的答案分析</li>
                <li><b>多种格式导出</b> - 支持PDF、Word等格式</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div style="background-color: #e6f3ff; padding: 20px; border-radius: 10px; height: 100%;">
            <h3 style="color: #0066cc; margin-bottom: 15px;">📚 使用指南</h3>
            <ol style="color: #444; padding-left: 20px;">
                <li><b>上传文件</b> - 提供题目、学生答案和评分标准</li>
                <li><b>选择语言</b> - 中文或英文评分反馈</li>
                <li><b>调整参数</b> - 根据需求设置评分严格度</li>
                <li><b>获取结果</b> - AI将提供详细分析和评分</li>
                <li><b>导出保存</b> - 将结果保存为需要的格式</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)

    # 添加间距
    st.markdown("<br>", unsafe_allow_html=True)

    # 快速功能导航
    st.markdown(
        "<h3 style='color: #444; margin-bottom: 15px; text-align: center;'>🚀 快速导航</h3>",
        unsafe_allow_html=True)

    # 三栏布局用于功能导航
    btn_col1, btn_col2, btn_col3 = st.columns(3)

    with btn_col1:
        if st.button("🖊️ AI 批改", use_container_width=True):
            if st.session_state.logged_in:
                st.session_state.page = "ai_correction"
                st.session_state.ai_correction_step = 1
            else:
                st.session_state.show_login_modal = True
                st.session_state.next_page_after_login = "ai_correction"
            st.session_state.needs_rerun = True

    with btn_col2:
        if st.button("📂 历史记录", use_container_width=True):
            if st.session_state.logged_in:
                st.session_state.page = "file_management"
            else:
                st.session_state.show_login_modal = True
                st.session_state.next_page_after_login = "file_management"
            st.session_state.needs_rerun = True

    with btn_col3:
        if st.button("👤 账户管理", use_container_width=True):
            if st.session_state.logged_in:
                st.session_state.page = "account_management"
            else:
                st.session_state.show_login_modal = True
                st.session_state.next_page_after_login = "account_management"
            st.session_state.needs_rerun = True

    # 添加页脚
    st.markdown("<br><br>", unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align: center; color: #888; font-size: 14px; margin-top: 50px;">
        <p>AI Guru © 2023 | 智能批改系统</p>
    </div>
    """, unsafe_allow_html=True)


def account_management_page():
    """Account management page with user profile and settings"""
    st.title("👤 Account Management")

    if not st.session_state.logged_in:
        st.warning("Please log in to access account management features.")
        if st.button("Log in now"):
            st.session_state.show_login_modal = True
            st.session_state.needs_rerun = True
        return

    # 用户信息部分
    st.subheader("🧑‍💼 User Profile")
    user_data = read_user_data()
    current_user_data = user_data.get(st.session_state.current_user, {})

    # 显示用户基本信息
    with st.container():
        st.markdown(f"""
        <div style="background-color: #f8f9fa; padding: 20px; border-radius: 10px; margin-bottom: 20px;">
            <table style="width: 100%;">
                <tr>
                    <td style="width: 40%; padding: 8px;"><b>Username:</b></td>
                    <td style="padding: 8px;">{st.session_state.current_user}</td>
                </tr>
                <tr>
                    <td style="padding: 8px;"><b>Email:</b></td>
                    <td style="padding: 8px;">{current_user_data.get('email', 'Not provided')}</td>
                </tr>
                <tr>
                    <td style="padding: 8px;"><b>Account Created:</b></td>
                    <td style="padding: 8px;">{current_user_data.get('created_at', 'Unknown')}</td>
                </tr>
                <tr>
                    <td style="padding: 8px;"><b>Correction Records:</b></td>
                    <td style="padding: 8px;">{len(current_user_data.get('records', []))}</td>
                </tr>
            </table>
        </div>
        """, unsafe_allow_html=True)

    # 账户安全
    st.subheader("🔒 Account Security")

    # 修改密码
    with st.expander("Change Password"):
        with st.form("change_password_form"):
            current_password = st.text_input(
                "Current Password", type="password")
            new_password = st.text_input("New Password", type="password")
            confirm_new_password = st.text_input(
                "Confirm New Password", type="password")

            password_submitted = st.form_submit_button("Update Password")

            if password_submitted:
                if not current_password or not new_password or not confirm_new_password:
                    st.error("All fields are required.")
                elif new_password != confirm_new_password:
                    st.error("New passwords do not match.")
                else:
                    # 检查当前密码是否正确
                    current_hash = current_user_data.get('password')

                    # 对于测试账户的特殊处理
                    if st.session_state.current_user in TEST_ACCOUNTS:
                        if current_password == TEST_ACCOUNTS[st.session_state.current_user]['password']:
                            # 更新密码
                            current_user_data['password'] = hash_password(
                                new_password)
                            user_data[st.session_state.current_user] = current_user_data
                            save_user_data(user_data)
                            st.success("Password updated successfully!")
                        else:
                            st.error("Current password is incorrect.")
                    else:
                        # 对于普通用户
                        if current_hash == hash_password(current_password):
                            # 更新密码
                            current_user_data['password'] = hash_password(
                                new_password)
                            user_data[st.session_state.current_user] = current_user_data
                            save_user_data(user_data)
                            st.success("Password updated successfully!")
                        else:
                            st.error("Current password is incorrect.")

    # 更新电子邮件
    with st.expander("Update Email"):
        with st.form("update_email_form"):
            current_email = current_user_data.get('email', '')
            st.write(
                f"Current email: {
                    current_email if current_email else 'Not set'}")

            new_email = st.text_input("New Email Address")

            email_submitted = st.form_submit_button("Update Email")

            if email_submitted:
                if not new_email:
                    st.error("Please enter a valid email address.")
                else:
                    # 更新邮箱
                    current_user_data['email'] = new_email
                    user_data[st.session_state.current_user] = current_user_data
                    save_user_data(user_data)
                    st.success("Email updated successfully!")

    # 页脚
    st.markdown("""
    <div style="text-align: center; margin-top: 30px; padding-top: 15px; border-top: 1px solid #eee; color: #666; font-size: 0.8em;">
        AI Guru User Management • Last updated: 2023
    </div>
    """, unsafe_allow_html=True)

# Add these helper functions after the save_uploaded_file function

def detect_file_groups(user_files):
    """
    将相关的上传文件分组，以便在显示时将属于同一组的文件放在一起
    规则：
    1. 具有相同数字编号的文件分在同一组
    2. 无编号的文件如果是同一题型分在同一组
    3. 如果只有一组答案，默认所有文件为同一组
    
    参数:
        user_files: 字典，包含上传的文件信息
        
    返回:
        list: 文件组列表，每组是一个文件键的列表
    """
    if not user_files:
        return []
        
    # 如果只有几个文件，通常是同一组
    if len(user_files) <= 3:
        return [list(user_files.keys())]
        
    # 按编号分组
    numbered_groups = {}
    unnumbered_files = []
    
    for file_key in user_files.keys():
        # 查找文件名中的数字
        number_match = re.search(r'(\d+)', file_key)
        
        if number_match:
            question_number = number_match.group(1)
            if question_number not in numbered_groups:
                numbered_groups[question_number] = []
            numbered_groups[question_number].append(file_key)
        else:
            unnumbered_files.append(file_key)
    
    # 构建最终的文件组
    file_groups = list(numbered_groups.values())
    
    # 处理没有编号的文件
    if unnumbered_files:
        # 如果还没有任何分组，所有无编号文件作为一组
        if not file_groups:
            file_groups.append(unnumbered_files)
        else:
            # 根据文件类型分配到已有的组中，或者创建新组
            for file_key in unnumbered_files:
                # 检查是否能分配到已有组
                assigned = False
                
                # 尝试根据文件类型匹配
                file_type = None
                if "question" in file_key:
                    file_type = "question"
                elif "student_answer" in file_key:
                    file_type = "student_answer"
                elif "marking_scheme" in file_key:
                    file_type = "marking_scheme"
                
                if file_type:
                    # 查找缺少此类型的组
                    for group in file_groups:
                        has_this_type = False
                        for existing_file in group:
                            if file_type in existing_file:
                                has_this_type = True
                                break
                        
                        if not has_this_type:
                            group.append(file_key)
                            assigned = True
                            break
                
                # 如果未分配，添加到第一个组
                if not assigned and file_groups:
                    file_groups[0].append(file_key)
    
    return file_groups

def parse_correction_blocks(correction_content):
    """
    解析批改内容，将其分成不同的批改块
    每个块通常对应一道题的批改
    
    参数:
        correction_content: 批改内容文本
        
    返回:
        list: 批改块列表，每块对应一道题的批改
    """
    if not correction_content:
        return []
        
    # 去除"学生答案批改如下:"开头
    content = re.sub(r'^学生答案批改如下:\s*', '', correction_content)
    
    # 查找题目标记，通常是"第X题"或"题目X"格式
    question_markers = re.findall(r'(第\s*\d+\s*题|题目\s*\d+)', content)
    
    # 如果没有找到题目标记，返回整个内容作为一个块
    if not question_markers:
        return [content]
        
    # 用题目标记分割内容
    blocks = []
    for i, marker in enumerate(question_markers):
        # 构建分割的正则表达式
        if i < len(question_markers) - 1:
            next_marker = question_markers[i + 1]
            pattern = f'({re.escape(marker)}.+?)(?={re.escape(next_marker)})'
        else:
            pattern = f'({re.escape(marker)}.+)'
            
        # 查找匹配的块
        match = re.search(pattern, content, re.DOTALL)
        if match:
            blocks.append(match.group(1).strip())
    
    # 如果没有成功分割，返回整个内容
    if not blocks:
        blocks = [content]
        
    return blocks

if __name__ == "__main__":
    main()
