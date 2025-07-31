# 修复版本的API调用模块 - 完整版
import base64
import requests  
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    openai = None
import re
from pathlib import Path
import json
import os
import time
import logging
from typing import Dict, List, Tuple, Any, Optional, Union
from dataclasses import dataclass
import contextlib
import io
from PIL import Image

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('logs/api_debug.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 强制使用简化版提示词
USE_SIMPLIFIED_PROMPTS = True
try:
    import prompts_simplified as prompts_module
    logger.info("使用简化版提示词系统")
except ImportError:
    try:
        from . import prompts_simplified as prompts_module
        logger.info("使用简化版提示词系统")
    except ImportError:
        logger.error("简化版提示词模块未找到")
        raise ImportError("无法导入简化版提示词模块")

# 导入PDF错误抑制工具
try:
    from pdf_utils import SuppressOutput, safe_pdf_processing
    PDF_UTILS_AVAILABLE = True
except ImportError:
    PDF_UTILS_AVAILABLE = False

# 抑制MuPDF错误输出
import warnings
warnings.filterwarnings("ignore")
os.environ['MUPDF_QUIET'] = '1'

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

@dataclass
class APIConfig:
    """API配置类 - 增强版"""
    api_key: str = ""
    base_url: str = "https://openrouter.ai/api/v1"
    
    # 多模型支持 - 按优先级排序
    available_models: list = None
    current_model_index: int = 0
    
    max_tokens: int = 50000
    temperature: float = 0.7
    max_retries: int = 3
    retry_delay: float = 1.0
    timeout: int = 120
    
    def __post_init__(self):
        """初始化后处理，从环境变量设置API密钥"""
        # 初始化可用模型列表
        if self.available_models is None:
            self.available_models = [
                "google/gemini-2.5-flash-lite-preview-06-17",  # 原始模型
                "google/gemini-2.5-flash",                    # 备用模型1
                "google/gemini-2.5-pro",                          # 备用模型2
                "anthropic/claude-3-haiku",                   # 备用模型3
                "meta-llama/llama-3-8b-instruct:free",       # 免费模型
                "microsoft/wizardlm-2-8x22b:free",           # 免费模型2
                "gryphe/mythomist-7b:free"                    # 免费模型3
            ]
        
        env_key = os.getenv('OPENROUTER_API_KEY') or os.getenv('OPENAI_API_KEY')
        if env_key:
            self.api_key = env_key
            return
        
        env_file_path = Path('.env')
        if env_file_path.exists():
            try:
                with open(env_file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith('#') and '=' in line:
                            key, value = line.split('=', 1)
                            if key.strip() == 'OPENROUTER_API_KEY' and value.strip():
                                self.api_key = value.strip()
                                return
            except Exception as e:
                logger.warning(f"读取.env文件失败: {e}")
        
        if not self.api_key:
            self.api_key = "请在此处输入您的新API密钥"
    
    @property
    def model(self) -> str:
        """获取当前使用的模型"""
        if self.current_model_index < len(self.available_models):
            return self.available_models[self.current_model_index]
        return self.available_models[0]  # 如果索引超出范围，返回第一个模型
    
    def switch_to_next_model(self) -> bool:
        """切换到下一个可用模型"""
        if self.current_model_index < len(self.available_models) - 1:
            self.current_model_index += 1
            logger.info(f"切换到备用模型: {self.model}")
            return True
        return False
    
    def reset_model(self):
        """重置到第一个模型"""
        self.current_model_index = 0
        logger.info(f"重置为主要模型: {self.model}")
    
    def is_valid(self) -> bool:
        """检查API配置是否有效"""
        return bool(self.api_key and self.api_key.startswith(('sk-', 'or-')))
    
    def get_status(self) -> dict:
        """获取配置状态信息"""
        api_key_source = "default"
        if os.getenv('OPENROUTER_API_KEY') or os.getenv('OPENAI_API_KEY'):
            api_key_source = "environment"
        elif Path('.env').exists():
            try:
                with open('.env', 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith('#') and '=' in line:
                            key, value = line.split('=', 1)
                            if key.strip() == 'OPENROUTER_API_KEY' and value.strip() and value.strip() != 'your_api_key_here':
                                api_key_source = ".env file"
                                break
            except:
                pass
        
        return {
            "api_key_configured": bool(self.api_key and self.api_key != "请在此处输入您的新API密钥"),
            "api_key_source": api_key_source,
            "base_url": self.base_url,
            "current_model": self.model,
            "available_models": self.available_models,
            "model_index": self.current_model_index,
            "is_valid": self.is_valid()
        }

# 全局配置实例
api_config = APIConfig()

def img_to_base64(image_path, max_size_mb=4):
    """将图片文件转换为base64编码，支持自动压缩"""
    import io
    import os
    from pathlib import Path
    from PIL import Image
    
    # 检查文件类型
    if isinstance(image_path, str):
        file_path = Path(image_path)
        if file_path.exists():
            file_ext = file_path.suffix.lower()
            if file_ext in ['.txt', '.md', '.doc', '.docx', '.rtf']:
                raise ValueError(f"文件 {image_path} 是文本文件，不能作为图像处理")
    
    # 处理URL
    if isinstance(image_path, str) and image_path.startswith(('http://', 'https://')):
        response = requests.get(image_path)
        response.raise_for_status()  
        image_data = response.content
    # 处理Streamlit上传的文件对象
    elif hasattr(image_path, 'read') and callable(image_path.read):
        try:
            if hasattr(image_path, 'tell') and callable(image_path.tell):
                current_position = image_path.tell()
            else:
                current_position = 0
            image_data = image_path.read()
            if hasattr(image_path, 'seek') and callable(image_path.seek):
                image_path.seek(current_position)
        except Exception as e:
            raise Exception(f"Failed to read uploaded file: {str(e)}")
    # 处理本地文件路径
    elif isinstance(image_path, str):
        file_size_mb = os.path.getsize(image_path) / (1024 * 1024)
        
        if file_size_mb <= max_size_mb:
            with open(image_path, "rb") as image_file:
                image_data = image_file.read()
        else:
            logging.info(f"图片文件过大 ({file_size_mb:.2f}MB)，正在压缩...")
            img = Image.open(image_path)
            
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            
            quality = 85
            max_dimension = 1920
            
            if max(img.size) > max_dimension:
                ratio = max_dimension / max(img.size)
                new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)
            
            while quality > 20:
                buffer = io.BytesIO()
                img.save(buffer, format='JPEG', quality=quality, optimize=True)
                compressed_size_mb = len(buffer.getvalue()) / (1024 * 1024)
                
                if compressed_size_mb <= max_size_mb:
                    logging.info(f"压缩完成: {file_size_mb:.2f}MB -> {compressed_size_mb:.2f}MB (质量: {quality})")
                    image_data = buffer.getvalue()
                    break
                
                quality -= 10
            else:
                while max_dimension > 800:
                    max_dimension -= 200
                    ratio = max_dimension / max(img.size)
                    new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                    img_resized = img.resize(new_size, Image.Resampling.LANCZOS)
                    
                    buffer = io.BytesIO()
                    img_resized.save(buffer, format='JPEG', quality=70, optimize=True)
                    compressed_size_mb = len(buffer.getvalue()) / (1024 * 1024)
                    
                    if compressed_size_mb <= max_size_mb:
                        logging.info(f"缩放压缩完成: {file_size_mb:.2f}MB -> {compressed_size_mb:.2f}MB (尺寸: {new_size})")
                        image_data = buffer.getvalue()
                        break
                else:
                    buffer = io.BytesIO()
                    img.save(buffer, format='JPEG', quality=50, optimize=True)
                    final_size_mb = len(buffer.getvalue()) / (1024 * 1024)
                    logging.info(f"最终压缩: {file_size_mb:.2f}MB -> {final_size_mb:.2f}MB")
                    image_data = buffer.getvalue()
    else:
        raise Exception(f"Unsupported image source type: {type(image_path)}")
        
    return base64.b64encode(image_data).decode('utf-8')

def get_file_type(file_path):
    """检测文件类型，返回文件类型分类"""
    if isinstance(file_path, str):
        file_ext = Path(file_path).suffix.lower()
        if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']:
            return 'image'
        elif file_ext == '.pdf':
            return 'pdf'
        elif file_ext in ['.doc', '.docx']:
            return 'word'
        elif file_ext in ['.txt', '.md', '.rtf', '.csv']:
            return 'text'
        elif file_ext in ['.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml']:
            return 'text'
    return 'unknown'

def read_pdf_file(file_path):
    """读取PDF文件内容，支持多种PDF库作为备用方案"""
    
    # 方法1: 尝试使用 pdfplumber（推荐，稳定性好）
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                return text.strip()
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"pdfplumber读取PDF失败: {str(e)}")
    
    # 方法2: 尝试使用 PyPDF2
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                return text.strip()
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyPDF2读取PDF失败: {str(e)}")
    
    # 方法3: 尝试使用 PyMuPDF（如果可用）
    try:
        import fitz
        doc = fitz.open(file_path)
        text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        doc.close()
        if text.strip():
            return text.strip()
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"PyMuPDF读取PDF失败: {str(e)}")
    
    # 如果所有方法都失败，返回错误信息
    return f"[PDF文件读取失败: {Path(file_path).name}] - 所有PDF处理库都无法读取此文件，请确保PDF文件格式正确且未损坏"

def read_word_file(file_path):
    """读取Word文档内容"""
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except ImportError:
        return f"[Word文档: {Path(file_path).name}] - 需要安装python-docx库来读取Word文档"
    except Exception as e:
        return f"[Word文档读取失败: {Path(file_path).name}] - 错误: {str(e)}"

def read_text_file(file_path):
    """读取文本文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

def process_file_content(file_path):
    """根据文件类型处理文件内容"""
    file_type = get_file_type(file_path)
    file_name = Path(file_path).name
    
    # 增加文件类型标识，帮助AI更好地识别
    file_category = ""
    if 'QUESTION' in file_name.upper():
        file_category = "[题目文件]"
    elif 'ANSWER' in file_name.upper():
        file_category = "[学生答案]"
    elif 'MARKING' in file_name.upper():
        file_category = "[批改标准]"
    
    try:
        if file_type == 'image':
            logger.info(f"处理图片文件 {file_category}: {file_name}")
            return 'image', file_path
        elif file_type == 'pdf':
            logger.info(f"处理PDF文件 {file_category}: {file_name}")
            
            # 1. 尝试文本提取
            try:
                pdf_text = read_pdf_file(file_path)
                # Check if text extraction was successful and content is substantial
                if pdf_text and not pdf_text.startswith("[PDF文件读取失败") and len(pdf_text) > 100:
                    logger.info(f"PDF文本提取成功: {file_name}")
                    return 'text', f"{file_category} [PDF文本内容: {file_name}]\n{pdf_text}"
                else:
                    logger.warning(f"PDF文本提取失败或内容过少，将尝试图像转换: {file_name}")
            except Exception as e:
                logger.error(f"PDF文本提取时发生异常: {e}")

            # 2. 如果文本提取失败或内容不足，则进行图像转换
            try:
                file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                if file_size_mb > 50:
                    return 'error', f"[PDF文件过大: {file_name}] - 文件大小 {file_size_mb:.1f}MB 超过限制"
                
                # 直接返回PDF路径，强制以图像模式处理
                logger.info(f"PDF文件处理完成: {file_name}, 共14页")
                return 'pdf', file_path
            except Exception as e:
                logger.warning(f"PDF处理异常: {str(e)}")
                return 'error', f"[PDF处理失败: {file_name}] - {str(e)}"
        elif file_type == 'word':
            content = read_word_file(file_path)
            logger.info(f"处理Word文件 {file_category}: {file_name}")
            return 'text', f"{file_category} [Word文档: {file_name}]\n{content}"
        elif file_type == 'text':
            content = read_text_file(file_path)
            logger.info(f"处理文本文件 {file_category}: {file_name}")
            return 'text', f"{file_category} [文本文件: {file_name}]\n{content}"
        else:
            try:
                content = read_text_file(file_path)
                logger.info(f"处理未知类型文件 {file_category}: {file_name}")
                return 'text', f"{file_category} [文件: {file_name}]\n{content}"
            except:
                return 'error', f"[不支持的文件类型: {file_name}] - 无法处理此文件"
    except Exception as e:
        return 'error', f"[文件处理错误: {file_name}] - {str(e)}"

def validate_pdf_file(pdf_path):
    """验证PDF文件的有效性和可读性"""
    validation_result = {
        'is_valid': False,
        'error_message': '',
        'file_info': {},
        'suggestions': []
    }
    
    try:
        # 检查文件是否存在
        if not os.path.exists(pdf_path):
            validation_result['error_message'] = f"文件不存在: {pdf_path}"
            validation_result['suggestions'].append("请检查文件路径是否正确")
            return validation_result
        
        # 检查文件大小
        file_size = os.path.getsize(pdf_path)
        validation_result['file_info']['size_bytes'] = file_size
        validation_result['file_info']['size_mb'] = file_size / (1024 * 1024)
        
        if file_size == 0:
            validation_result['error_message'] = "PDF文件为空"
            validation_result['suggestions'].append("请检查文件是否正确上传")
            return validation_result
        
        if file_size > 100 * 1024 * 1024:  # 100MB
            validation_result['error_message'] = f"PDF文件过大 ({file_size / (1024*1024):.1f}MB)"
            validation_result['suggestions'].append("请尝试压缩PDF文件或分割成多个小文件")
            return validation_result
        
        # 检查文件扩展名
        if not pdf_path.lower().endswith('.pdf'):
            validation_result['error_message'] = "文件扩展名不是.pdf"
            validation_result['suggestions'].append("请确保文件是PDF格式")
            return validation_result
        
        # 检查文件头部（PDF魔数）
        try:
            with open(pdf_path, 'rb') as f:
                header = f.read(8)
                if not header.startswith(b'%PDF-'):
                    validation_result['error_message'] = "文件不是有效的PDF格式（缺少PDF头部标识）"
                    validation_result['suggestions'].append("文件可能已损坏或不是真正的PDF文件")
                    return validation_result
                
                # 提取PDF版本
                try:
                    version_line = header.decode('ascii', errors='ignore')
                    if '%PDF-' in version_line:
                        version = version_line.split('%PDF-')[1][:3]
                        validation_result['file_info']['pdf_version'] = version
                except:
                    pass
        except Exception as e:
            validation_result['error_message'] = f"无法读取文件头部: {str(e)}"
            validation_result['suggestions'].append("文件可能已损坏或被其他程序占用")
            return validation_result
        
        # 尝试用PDF库验证文件结构
        pdf_library_results = []
        
        # 测试PyMuPDF
        try:
            import fitz
            doc = fitz.open(pdf_path)
            page_count = doc.page_count
            is_encrypted = doc.is_encrypted
            doc.close()
            
            pdf_library_results.append({
                'library': 'PyMuPDF',
                'success': True,
                'page_count': page_count,
                'is_encrypted': is_encrypted
            })
            
            validation_result['file_info']['page_count'] = page_count
            validation_result['file_info']['is_encrypted'] = is_encrypted
            
            if is_encrypted:
                validation_result['error_message'] = "PDF文件已加密"
                validation_result['suggestions'].append("请提供未加密的PDF文件或提供密码")
                return validation_result
            
            if page_count == 0:
                validation_result['error_message'] = "PDF文件没有页面"
                validation_result['suggestions'].append("文件可能已损坏")
                return validation_result
                
        except ImportError:
            pdf_library_results.append({
                'library': 'PyMuPDF',
                'success': False,
                'error': 'Library not installed'
            })
        except Exception as e:
            pdf_library_results.append({
                'library': 'PyMuPDF',
                'success': False,
                'error': str(e)
            })
        
        # 测试pypdfium2
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(pdf_path)
            page_count = len(pdf)
            pdf.close()
            
            pdf_library_results.append({
                'library': 'pypdfium2',
                'success': True,
                'page_count': page_count
            })
            
            if not validation_result['file_info'].get('page_count'):
                validation_result['file_info']['page_count'] = page_count
                
        except ImportError:
            pdf_library_results.append({
                'library': 'pypdfium2',
                'success': False,
                'error': 'Library not installed'
            })
        except Exception as e:
            pdf_library_results.append({
                'library': 'pypdfium2',
                'success': False,
                'error': str(e)
            })
        
        validation_result['file_info']['library_tests'] = pdf_library_results
        
        # 检查是否至少有一个库能成功处理
        successful_libraries = [r for r in pdf_library_results if r['success']]
        if not successful_libraries:
            validation_result['error_message'] = "所有PDF处理库都无法打开此文件"
            validation_result['suggestions'].extend([
                "文件可能已损坏",
                "请尝试用其他PDF阅读器打开文件验证",
                "如果文件正常，请检查是否安装了PyMuPDF或pypdfium2库"
            ])
            return validation_result
        
        # 如果所有检查都通过
        validation_result['is_valid'] = True
        validation_result['error_message'] = ''
        
        return validation_result
        
    except Exception as e:
        validation_result['error_message'] = f"PDF验证过程中发生错误: {str(e)}"
        validation_result['suggestions'].append("请联系技术支持")
        return validation_result

def pdf_pages_to_base64_images(pdf_path, zoom=2.0):
    """将 PDF 每页转换为 Base64 编码的图像数据列表，支持多种方法 - 增强错误诊断版"""
    
    base64_images = []
    error_details = []
    
    # 首先进行PDF文件验证
    logger.info(f"开始验证PDF文件: {os.path.basename(pdf_path)}")
    validation_result = validate_pdf_file(pdf_path)
    
    if not validation_result['is_valid']:
        error_msg = f"PDF文件验证失败: {validation_result['error_message']}"
        logger.error(error_msg)
        logger.error("建议解决方案:")
        for suggestion in validation_result['suggestions']:
            logger.error(f"  - {suggestion}")
        
        # 记录文件信息用于调试
        if validation_result['file_info']:
            logger.error("文件信息:")
            for key, value in validation_result['file_info'].items():
                logger.error(f"  {key}: {value}")
        
        return []
    
    # 记录验证成功的文件信息
    file_info = validation_result['file_info']
    logger.info(f"PDF文件验证成功: {os.path.basename(pdf_path)}")
    logger.info(f"  - 文件大小: {file_info.get('size_mb', 0):.1f}MB")
    logger.info(f"  - 页面数量: {file_info.get('page_count', 'unknown')}")
    logger.info(f"  - PDF版本: {file_info.get('pdf_version', 'unknown')}")
    logger.info(f"  - 是否加密: {file_info.get('is_encrypted', 'unknown')}")
    
    # 方法1: 尝试使用 PyMuPDF (fitz)
    fitz_available = False
    try:
        import fitz
        fitz_available = True
        logger.info(f"PyMuPDF可用，版本: {fitz.version}")
        
        suppress_context = SuppressOutput() if PDF_UTILS_AVAILABLE else contextlib.nullcontext()
        
        with suppress_context:
            try:
                doc = fitz.open(pdf_path)
                logger.info(f"PDF文件打开成功: {pdf_path}")
                
                if doc.is_encrypted:
                    logger.warning(f"PDF文件 {pdf_path} 已加密，尝试解密")
                    if not doc.authenticate(""):
                        error_msg = f"PDF文件已加密且无法解密: {pdf_path}"
                        logger.error(error_msg)
                        error_details.append(error_msg)
                        doc.close()
                    else:
                        logger.info("PDF解密成功")
                
                if not doc.is_encrypted or doc.authenticate(""):
                    if doc.page_count == 0:
                        error_msg = f"PDF文件没有页面: {pdf_path}"
                        logger.warning(error_msg)
                        error_details.append(error_msg)
                        doc.close()
                    else:
                        max_pages = min(doc.page_count, 20)  # 限制最多处理20页
                        logger.info(f"开始处理PDF，共{doc.page_count}页，处理前{max_pages}页")
                        
                        for page_num in range(max_pages):
                            try:
                                with contextlib.redirect_stderr(open(os.devnull, 'w')):
                                    page = doc.load_page(page_num)
                                    matrix = fitz.Matrix(zoom, zoom)
                                    pix = page.get_pixmap(matrix=matrix, alpha=False)
                                    img_data = pix.tobytes("png")
                                    
                                    # 检查图像数据大小
                                    if len(img_data) > 5 * 1024 * 1024:  # 5MB
                                        logger.info(f"第{page_num + 1}页图像过大({len(img_data)/(1024*1024):.1f}MB)，进行压缩")
                                        img = Image.open(io.BytesIO(img_data))
                                        if img.mode in ('RGBA', 'LA', 'P'):
                                            img = img.convert('RGB')
                                        
                                        max_size = 1600
                                        if max(img.size) > max_size:
                                            ratio = max_size / max(img.size)
                                            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
                                            img = img.resize(new_size, Image.Resampling.LANCZOS)
                                        
                                        buffer = io.BytesIO()
                                        img.save(buffer, format='JPEG', quality=80, optimize=True)
                                        img_data = buffer.getvalue()
                                        logger.info(f"压缩后大小: {len(img_data)/(1024*1024):.1f}MB")
                                    
                                    base64_str = base64.b64encode(img_data).decode("utf-8")
                                    base64_images.append(base64_str)
                                    pix = None
                                    
                                    logger.debug(f"成功处理第{page_num + 1}页")
                                    
                            except Exception as e:
                                error_msg = f"PyMuPDF处理PDF第{page_num + 1}页时出错: {str(e)}"
                                logger.warning(error_msg)
                                error_details.append(error_msg)
                                continue
                        
                        doc.close()
                        
                        if base64_images:
                            logger.info(f"PyMuPDF成功处理PDF文件 {pdf_path}，共{len(base64_images)}页")
                            return base64_images
                        else:
                            error_msg = "PyMuPDF处理完成但未生成任何图像"
                            logger.warning(error_msg)
                            error_details.append(error_msg)
                            
            except Exception as doc_error:
                error_msg = f"PyMuPDF打开PDF文件失败: {str(doc_error)}"
                logger.warning(error_msg)
                error_details.append(error_msg)
            
    except ImportError:
        error_msg = "PyMuPDF未安装"
        logger.warning(error_msg)
        error_details.append(error_msg)
    except Exception as e:
        error_msg = f"PyMuPDF处理PDF失败: {str(e)}"
        logger.warning(error_msg)
        error_details.append(error_msg)
    
    # 方法2: 使用 pypdfium2 (通过 pdfplumber 依赖安装)
    pypdfium2_available = False
    try:
        import pypdfium2 as pdfium
        pypdfium2_available = True
        logger.info("pypdfium2可用")
        
        try:
            pdf = pdfium.PdfDocument(pdf_path)
            max_pages = min(len(pdf), 20)  # 限制最多处理20页
            logger.info(f"pypdfium2打开PDF成功，共{len(pdf)}页，处理前{max_pages}页")
            
            for page_num in range(max_pages):
                try:
                    page = pdf.get_page(page_num)
                    # 渲染页面为图像
                    bitmap = page.render(scale=zoom, rotation=0)
                    pil_image = bitmap.to_pil()
                    
                    # 转换为RGB格式
                    if pil_image.mode != 'RGB':
                        pil_image = pil_image.convert('RGB')
                    
                    # 压缩图像如果太大
                    if max(pil_image.size) > 1600:
                        ratio = 1600 / max(pil_image.size)
                        new_size = (int(pil_image.size[0] * ratio), int(pil_image.size[1] * ratio))
                        pil_image = pil_image.resize(new_size, Image.Resampling.LANCZOS)
                    
                    # 转换为base64
                    buffer = io.BytesIO()
                    pil_image.save(buffer, format='JPEG', quality=80, optimize=True)
                    img_data = buffer.getvalue()
                    base64_str = base64.b64encode(img_data).decode("utf-8")
                    base64_images.append(base64_str)
                    
                    page.close()
                    logger.debug(f"pypdfium2成功处理第{page_num + 1}页")
                    
                except Exception as e:
                    error_msg = f"pypdfium2处理PDF第{page_num + 1}页时出错: {str(e)}"
                    logger.warning(error_msg)
                    error_details.append(error_msg)
                    continue
            
            pdf.close()
            
            if base64_images:
                logger.info(f"pypdfium2成功处理PDF文件 {pdf_path}，共{len(base64_images)}页")
                return base64_images
            else:
                error_msg = "pypdfium2处理完成但未生成任何图像"
                logger.warning(error_msg)
                error_details.append(error_msg)
                
        except Exception as pdf_error:
            error_msg = f"pypdfium2打开PDF文件失败: {str(pdf_error)}"
            logger.warning(error_msg)
            error_details.append(error_msg)
            
    except ImportError:
        error_msg = "pypdfium2未安装"
        logger.warning(error_msg)
        error_details.append(error_msg)
    except Exception as e:
        error_msg = f"pypdfium2处理PDF失败: {str(e)}"
        logger.warning(error_msg)
        error_details.append(error_msg)
    
    # 如果所有方法都失败，提供详细的错误信息
    if not base64_images:
        logger.error(f"所有PDF处理方法都失败，无法生成预览: {pdf_path}")
        logger.error("错误详情:")
        for i, error in enumerate(error_details, 1):
            logger.error(f"  {i}. {error}")
        
        # 提供诊断信息
        diagnostic_info = []
        diagnostic_info.append(f"文件路径: {pdf_path}")
        diagnostic_info.append(f"文件存在: {os.path.exists(pdf_path)}")
        if os.path.exists(pdf_path):
            diagnostic_info.append(f"文件大小: {os.path.getsize(pdf_path)} 字节")
        diagnostic_info.append(f"PyMuPDF可用: {fitz_available}")
        diagnostic_info.append(f"pypdfium2可用: {pypdfium2_available}")
        
        logger.error("诊断信息:")
        for info in diagnostic_info:
            logger.error(f"  - {info}")
        
        return []
    
    return base64_images

def create_batch_grading_prompt(batch_number, total_batches, current_range, system_message):
    """创建分批批改的提示词"""
    batch_prompt = f"""
🔢【分批批改 - 第{batch_number}批/共{total_batches}批】🔢

📋 当前批次范围：第{current_range[0]}-{current_range[1]}题
⚠️ 只关注当前批次的题目，不要涉及其他题目

{system_message}

🎯【当前批次专项要求】🎯
1. 仅批改第{current_range[0]}-{current_range[1]}题
2. 如果学生答案中没有某题，标记为"未作答"并跳过
3. 完成当前批次后立即停止，不要继续其他题目
4. 每题批改后检查是否出现重复内容

⏸️【批次完成检查】⏸️
当前批次预期完成第{current_range[0]}-{current_range[1]}题：
□ 是否只处理了当前批次范围内的题目？
□ 是否跳过了学生未作答的题目？
□ 是否避免了循环和重复？
□ 是否准备好进入下一批次？

开始第{batch_number}批批改：
"""
    return batch_prompt

def split_grading_task(content_text, batch_size=10):
    """
    将大型批改任务分割成小批次
    
    Args:
        content_text: 要批改的内容
        batch_size: 每批次处理的题目数量
    
    Returns:
        list: 分批处理的配置列表
    """
    # 简单的题目数量估算（基于常见模式）
    import re
    
    # 寻找题目标识符
    question_patterns = [
        r'第\s*(\d+)\s*题',  # 第X题
        r'题\s*(\d+)',      # 题X
        r'Q\.?\s*(\d+)',    # Q.X 或 QX
        r'(\d+)\)',         # X)
        r'Problem\s*(\d+)', # Problem X
    ]
    
    question_numbers = set()
    for pattern in question_patterns:
        matches = re.findall(pattern, content_text, re.IGNORECASE)
        for match in matches:
            try:
                question_numbers.add(int(match))
            except ValueError:
                continue
    
    if not question_numbers:
        # 如果没有找到题目编号，假设有10道题
        question_numbers = set(range(1, 11))
    
    max_question = max(question_numbers) if question_numbers else 10
    
    # 创建批次配置
    batches = []
    for start in range(1, max_question + 1, batch_size):
        end = min(start + batch_size - 1, max_question)
        batches.append({
            'range': (start, end),
            'batch_number': len(batches) + 1,
            'total_batches': (max_question + batch_size - 1) // batch_size
        })
    
    return batches

def call_tongyiqianwen_api_batch(input_text: str, *input_contents, system_message: str = "", batch_size: int = 10) -> str:
    """
    分批调用API进行批改，避免循环和内存溢出
    
    Args:
        input_text: 输入文本
        input_contents: 输入内容（图片等）
        system_message: 系统消息
        batch_size: 每批次处理的题目数量
    
    Returns:
        str: 完整的批改结果
    """
    try:
        # 检查是否需要分批处理
        content_length = len(input_text)
        if content_length < 5000:  # 内容较少，不需要分批
            return call_tongyiqianwen_api(input_text, *input_contents, system_message=system_message)
        
        logger.info(f"内容较长({content_length}字符)，启动分批处理模式")
        
        # 分析题目范围
        batches = split_grading_task(input_text, batch_size)
        logger.info(f"计划分{len(batches)}批处理，每批{batch_size}题")
        
        all_results = []
        
        for i, batch_config in enumerate(batches):
            batch_number = batch_config['batch_number']
            total_batches = batch_config['total_batches']
            current_range = batch_config['range']
            
            logger.info(f"开始处理第{batch_number}批：第{current_range[0]}-{current_range[1]}题")
            
            # 创建当前批次的提示词
            batch_prompt = create_batch_grading_prompt(
                batch_number, total_batches, current_range, system_message
            )
            
            # 调用API处理当前批次
            try:
                batch_result = call_tongyiqianwen_api(
                    input_text, 
                    *input_contents, 
                    system_message=batch_prompt
                )
                
                if batch_result:
                    # 添加批次标识
                    batch_header = f"\n{'='*50}\n📋 第{batch_number}批批改结果 (第{current_range[0]}-{current_range[1]}题)\n{'='*50}\n"
                    all_results.append(batch_header + batch_result)
                    
                    logger.info(f"第{batch_number}批处理完成")
                else:
                    logger.warning(f"第{batch_number}批处理失败，跳过")
                    
            except Exception as e:
                logger.error(f"第{batch_number}批处理出错: {str(e)}")
                all_results.append(f"\n❌ 第{batch_number}批处理失败: {str(e)}\n")
                continue
        
        # 合并所有结果
        if all_results:
            final_result = "\n".join(all_results)
            
            # 添加总结
            summary = f"""
\n{'='*50}
📊 分批批改总结
{'='*50}
✅ 共处理 {len(batches)} 个批次
📋 预期题目范围: 第1-{batches[-1]['range'][1]}题
⚠️ 如有题目缺失，说明学生未作答
✨ 批改完成！
{'='*50}
"""
            final_result += summary
            
            logger.info("分批批改全部完成")
            return final_result
        else:
            logger.error("所有批次都处理失败")
            return "❌ 分批批改失败，所有批次都出现错误"
            
    except Exception as e:
        logger.error(f"分批批改系统出错: {str(e)}")
        # 降级到普通模式
        logger.info("降级到普通批改模式")
        return call_tongyiqianwen_api(input_text, *input_contents, system_message=system_message)

def call_tongyiqianwen_api(input_text: str, *input_contents, system_message: str = "") -> str:
    """调用API进行多类型文件处理"""
    from openai import OpenAI
    
    if not api_config.is_valid():
        error_msg = f"""
🚫 API配置错误

可能的解决方案：
1. 设置环境变量：OPENROUTER_API_KEY=your_api_key
2. 检查API密钥格式
3. 确认密钥有效性

当前配置状态：
{json.dumps(api_config.get_status(), ensure_ascii=False, indent=2)}"""
        logger.error("API配置无效")
        return error_msg
    
    try:
        client = OpenAI(
            api_key=api_config.api_key,
            base_url=api_config.base_url,
            timeout=api_config.timeout
        )
    except Exception as e:
        error_msg = f"❌ OpenAI客户端初始化失败: {str(e)}"
        logger.error(error_msg)
        return error_msg
    
    content = [{"type": "text", "text": input_text}]
    
    try:
        for single_content in input_contents:
            if (isinstance(single_content, tuple) and 
                len(single_content) == 2 and 
                all(isinstance(item, str) for item in single_content)):
                content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/{single_content[0]};base64,{single_content[1]}"
                    }
                })   
            elif os.path.isfile(single_content):
                logger.info(f"处理文件 [识别类型]: {os.path.basename(single_content)}")
                content_type, processed_content = process_file_content(single_content)            
                if content_type == 'text':
                    logger.info(f"文本文件处理完成: {os.path.basename(single_content)}, 长度: {len(processed_content)} 字符")
                    content.append({
                        "type": "text",
                        "text": processed_content
                    })
                elif content_type == 'image':
                    logger.info(f"图片文件处理开始: {os.path.basename(single_content)}")
                    base_64_image = img_to_base64(single_content)
                    logger.info(f"图片文件处理完成: {os.path.basename(single_content)}, Base64长度: {len(base_64_image)}")
                    content.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base_64_image}"
                        }
                    })    
                elif content_type == 'pdf':
                    logger.info(f"PDF文件处理开始: {os.path.basename(single_content)}")
                    base_64_images = pdf_pages_to_base64_images(single_content)
                    logger.info(f"PDF文件处理完成: {os.path.basename(single_content)}, 共{len(base_64_images)}页")
                    for i, base_64_image in enumerate(base_64_images):
                        content.append({
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base_64_image}"
                            }
                        })
                        logger.debug(f"PDF第{i+1}页已添加到内容中")
                else:
                    raise ValueError(f"The file {single_content} could not be processed.")
            else:
                content.append({
                    "type": "text",
                    "text": single_content
                })
    except Exception as e:
        error_msg = f"❌ 文件处理失败: {str(e)}"
        logger.error(error_msg)
        return error_msg

    for attempt in range(api_config.max_retries):
        try:
            final_message = []
            if system_message:
                final_message.append({"role": "system", "content": system_message})
            final_message.append({"role": "user", "content": content})
            
            # 记录API调用详细信息
            start_time = time.time()
            logger.info(f"API调用尝试 {attempt + 1}/{api_config.max_retries}")
            logger.info(f"发送的消息数量: {len(final_message)}")
            logger.info(f"使用的模型: {api_config.model}")
            
            # 记录输入内容统计
            total_text_length = sum(len(str(msg.get('content', ''))) for msg in final_message)
            logger.info(f"输入内容总长度: {total_text_length} 字符")
            
            response = client.chat.completions.create(
                model=api_config.model,
                messages=final_message,
                max_tokens=api_config.max_tokens,
                temperature=api_config.temperature
            )

            result = response.choices[0].message.content
            processing_time = time.time() - start_time
            
            # 记录API调用结果详细信息
            logger.info("=" * 80)
            logger.info("🔍 API调用完整返回结果:")
            logger.info(f"返回内容长度: {len(result) if result else 0}")
            logger.info(f"处理时间: {processing_time:.2f} 秒")
            logger.info(f"返回内容预览: {result[:200] if result else 'None'}...")
            logger.info(f"返回内容完整: {result if result else 'None'}")
            logger.info("=" * 80)
        
            if not result or not result.strip():
                logger.warning("API返回空结果")
                if attempt < api_config.max_retries - 1:
                    time.sleep(api_config.retry_delay)
                    continue
                else:
                    fallback_msg = "❌ API返回了空结果。可能的原因：文件内容无法识别或API服务暂时不可用。"
                    logger.error("所有重试都失败，返回fallback消息")
                    return fallback_msg
            
            logger.info("✅ API调用成功完成")
            return result
        
        except Exception as e:
            error_str = str(e)
            logger.error(f"API调用失败 (尝试 {attempt + 1}): {error_str}")
            
            if "timeout" in error_str.lower() or "timed out" in error_str.lower():
                timeout_error_msg = f"""❌ 请求超时错误
问题分析：网络连接超时、API服务器响应缓慢、处理的文件过大或过多
解决方案：检查网络连接、减少单次处理的文件数量、稍后重试
错误详情：{error_str}"""
                if attempt < api_config.max_retries - 1:
                    wait_time = api_config.retry_delay * (2 ** attempt)
                    logger.info(f"遇到超时错误，等待 {wait_time} 秒后重试")
                    time.sleep(wait_time)
                    continue
                else:
                    return timeout_error_msg
            
            if "401" in error_str or "Unauthorized" in error_str:
                auth_error_msg = f"""❌ 认证失败 (401 Unauthorized)
问题分析：API密钥无效或已过期、密钥格式错误、账户余额不足
解决方案：检查API密钥、更新API密钥、检查账户状态
当前使用的密钥来源：{api_config.get_status()['api_key_source']}
原始错误：{error_str}"""
                logger.error("认证失败")
                return auth_error_msg
            
            elif "429" in error_str or "rate_limit" in error_str.lower():
                # 尝试切换到备用模型
                if api_config.switch_to_next_model():
                    logger.info(f"遇到频率限制，切换到备用模型: {api_config.model}")
                    # 重新创建客户端
                    try:
                        client = OpenAI(
                            api_key=api_config.api_key,
                            base_url=api_config.base_url,
                            timeout=api_config.timeout
                        )
                        # 重置重试计数器，给新模型机会
                        attempt = 0
                        continue
                    except Exception as model_switch_error:
                        logger.error(f"切换模型失败: {model_switch_error}")
                
                # 如果没有更多模型可切换，则等待重试
                rate_limit_msg = f"❌ API调用频率限制，当前模型：{api_config.model}。错误：{error_str}"
                if attempt < api_config.max_retries - 1:
                    wait_time = api_config.retry_delay * (2 ** attempt)
                    logger.info(f"遇到频率限制，等待 {wait_time} 秒后重试")
                    time.sleep(wait_time)
                    continue
                else:
                    # 重置模型索引，为下次调用准备
                    api_config.reset_model()
                    return rate_limit_msg
            
            elif "500" in error_str or "502" in error_str or "503" in error_str or "504" in error_str:
                if "504" in error_str:
                    server_error_msg = f"""❌ 网关超时错误 (504 Gateway Timeout)
问题分析：API服务器响应超时、网络连接不稳定、服务器负载过高
解决方案：检查网络连接稳定性、稍后重试、考虑减少单次处理的文件数量
错误详情：{error_str}"""
                else:
                    server_error_msg = f"❌ 服务器错误，请稍后重试。错误：{error_str}"
                
                if attempt < api_config.max_retries - 1:
                    wait_time = api_config.retry_delay * (2 ** attempt)
                    logger.info(f"遇到服务器错误，等待 {wait_time} 秒后重试")
                    time.sleep(wait_time)
                    continue
                else:
                    return server_error_msg
            
            if attempt < api_config.max_retries - 1:
                time.sleep(api_config.retry_delay * (attempt + 1))
                continue
            else:
                error_msg = f"""❌ API调用失败 (所有重试已耗尽)
错误详情：{error_str}
可能的解决方案：检查网络连接、验证API密钥有效性、确认账户余额充足、稍后重试
配置信息：{json.dumps(api_config.get_status(), ensure_ascii=False, indent=2)}"""
                logger.error(error_msg)
                return error_msg

# 标准API调用函数
default_api = call_tongyiqianwen_api

# ===================== 结果类和装饰器 =====================
class GradingResult:
    """批改结果标准化类"""
    
    def __init__(self, success: bool = True, data: Any = None, error_message: str = "", processing_time: float = 0.0):
        self.success = success
        self.data = data
        self.error_message = error_message
        self.processing_time = processing_time
        self.timestamp = time.time()
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式，适合网站API返回"""
        return {
            "success": self.success,
            "data": self.data,
            "error_message": self.error_message,
            "processing_time": self.processing_time,
            "timestamp": self.timestamp
        }
    
    def to_json(self) -> str:
        """转换为JSON字符串"""
        return json.dumps(self.to_dict(), ensure_ascii=False, indent=2)

def safe_api_call(func):
    """装饰器：为API调用添加统一的错误处理和结果包装"""
    def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            logger.info(f"开始执行 {func.__name__}")
            result = func(*args, **kwargs)
            processing_time = time.time() - start_time
            logger.info(f"{func.__name__} 执行成功，耗时: {processing_time:.2f}秒")
            return GradingResult(success=True, data=result, processing_time=processing_time)
        except Exception as e:
            processing_time = time.time() - start_time
            error_msg = f"{func.__name__} 执行失败: {str(e)}"
            logger.error(error_msg)
            return GradingResult(success=False, error_message=error_msg, processing_time=processing_time)
    return wrapper

# ===================== 核心批改函数 =====================
def batch_correction_with_standard(marking_scheme_files: List[str], student_answer_files: List[str], 
                                  strictness_level: str = "中等", api=default_api, use_batch_processing: bool = True, batch_size: int = 10) -> dict:
    """批量批改 - 有批改标准模式"""
    try:
        marking_contents = []
        marking_file_info = []
        for i, file in enumerate(marking_scheme_files):
            content_type, content = process_file_content(file)
            if content_type == 'error':
                raise ValueError(f"处理批改标准文件失败: {content}")
            elif content_type == 'image' or content_type == 'pdf':
                marking_contents.append(file)
                marking_file_info.append(f"【批改方案文件 {i+1}】: {os.path.basename(file)}")
            else:
                marking_contents.append(content)
                marking_file_info.append(f"【批改方案文件 {i+1}】: {os.path.basename(file)}")
        
        student_contents = []
        student_file_info = []
        for i, file in enumerate(student_answer_files):
            content_type, content = process_file_content(file)
            if content_type == 'error':
                raise ValueError(f"处理学生答案文件失败: {content}")
            elif content_type == 'image' or content_type == 'pdf':
                student_contents.append(file)
                student_file_info.append(f"【学生作答文件 {i+1}】: {os.path.basename(file)}")
            else:
                student_contents.append(content)
                student_file_info.append(f"【学生作答文件 {i+1}】: {os.path.basename(file)}")
        
        prompt = prompts_module.get_complete_grading_prompt(file_info_list=[])
        
        api_args = []
        if marking_contents:
            api_args.append("=" * 50)
            api_args.append("批改方案文件（包含正确答案和评分标准）：")
            api_args.append("=" * 50)
            for i, (info, content) in enumerate(zip(marking_file_info, marking_contents)):
                api_args.append(f"\n{info}")
                api_args.append(content)
            api_args.append("\n" + "=" * 50)
        
        if student_contents:
            api_args.append("\n" + "=" * 50)
            api_args.append("学生作答文件（需要批改的答案）：")
            api_args.append("=" * 50)
            for i, (info, content) in enumerate(zip(student_file_info, student_contents)):
                api_args.append(f"\n{info}")
                api_args.append(content)
            api_args.append("\n" + "=" * 50)
        
        api_args.append("\n【批改指令】：")
        api_args.append(prompt)
        
        result = api(*api_args, system_message=prompts_module.get_core_grading_prompt())
        
        return {
            "correction_result": result,
            "has_separate_scheme": False
        }
        
    except Exception as e:
        error_msg = f"批改失败: {str(e)}"
        logger.error(error_msg)
        raise RuntimeError(error_msg) from e

def batch_correction_without_standard(question_files: List[str], student_answer_files: List[str], 
                                     strictness_level: str = "中等", api=default_api) -> dict:
    """批量批改 - 无批改标准模式（自动生成批改标准）"""
    try:
        logger.info("正在根据题目生成批改标准...")
        
        # 对于简化版，我们直接使用学生答案文件进行批改
        # 因为简化版假设MARKING_文件包含完整信息
        if not student_answer_files:
            raise ValueError("缺少学生答案文件")
        
        # 使用简化版批改逻辑
        file_info_list = []
        for file_path in student_answer_files:
            filename = os.path.basename(file_path)
            file_info = {
                'name': filename,
                'path': file_path,
                'type': get_file_type(file_path),
                'expected_category': 'answer'
            }
            file_info_list.append(file_info)
        
        prompt = prompts_module.get_complete_grading_prompt(file_info_list=file_info_list)
        
        api_args = []
        api_args.append("=" * 50)
        api_args.append("学生作答文件（需要批改的答案）：")
        api_args.append("=" * 50)
        
        for file_path in student_answer_files:
            api_args.append(f"\n文件：{os.path.basename(file_path)}")
            content_type, content = process_file_content(file_path)
            if content_type == 'error':
                raise ValueError(f"处理学生答案文件失败: {content}")
            elif content_type in ['image', 'pdf']:
                api_args.append(file_path)
            else:
                api_args.append(content)
        
        api_args.append("\n" + "=" * 50)
        api_args.append("\n【批改指令】：")
        api_args.append(prompt)
        
        result = api(*api_args, system_message=prompts_module.get_core_grading_prompt())
        
        return {
            "correction_result": result,
            "has_separate_scheme": False
        }
        
    except Exception as e:
        error_msg = f"批改失败: {str(e)}"
        logger.error(error_msg)
        raise RuntimeError(error_msg) from e

def intelligent_correction_with_files(question_files=None, answer_files=None, marking_scheme_files=None, 
                                    strictness_level="中等", mode="efficient"):
    """智能文件批改函数 - 简化版本"""
    if not answer_files:
        error_msg = "必须提供学生答案文件"
        logger.error(error_msg)
        return error_msg
    
    try:
        if marking_scheme_files:
            logger.info(f"使用批改标准模式 - 标准文件: {len(marking_scheme_files)}, 答案文件: {len(answer_files)}")
            result = batch_correction_with_standard(
                marking_scheme_files,
                answer_files,
                strictness_level=strictness_level,
                api=default_api
            )
        else:
            logger.info(f"使用自动生成批改标准模式 - 答案文件: {len(answer_files)}")
            files_for_scheme = question_files if question_files else answer_files
            result = batch_correction_without_standard(
                files_for_scheme,
                answer_files,
                strictness_level=strictness_level,
                api=default_api
            )
        
        return result
        
    except Exception as e:
        error_msg = f"智能批改失败: {str(e)}"
        logger.error(error_msg)
        return error_msg

def simplified_batch_correction(files_dict: dict, strictness_level: str = "严格", api=None) -> dict:
    """使用简化提示词的批量批改函数 - 修复版本"""
    if api is None:
        api = default_api
    
    try:
        answer_files = files_dict.get('answer_files', [])
        marking_files = files_dict.get('marking_files', [])
        question_files = files_dict.get('question_files', [])
        
        if not answer_files:
            return {
                "success": False,
                "error": "缺少学生答案文件",
                "message": "必须提供学生答案文件"
            }
        
        # 构建文件信息列表
        file_info_list = []
        all_files = [
            (answer_files, 'answer'),
            (marking_files, 'marking'),
            (question_files, 'question')
        ]
        
        for file_list, expected_category in all_files:
            for file_path in file_list:
                filename = os.path.basename(file_path)
                file_info = {
                    'name': filename,
                    'path': file_path,
                    'type': get_file_type(file_path),
                    'expected_category': expected_category
                }
                file_info_list.append(file_info)
        
        prompt = prompts_module.get_complete_grading_prompt(file_info_list=file_info_list)
        system_msg = prompts_module.get_core_grading_prompt()
        
        api_args = []
        
        # 检查文件完整性
        if not answer_files:
            return {
                "success": False,
                "error": "缺少学生答案文件",
                "message": "缺少学生答案文件"
            }
        
        # 处理文件
        if question_files:
            api_args.append("=" * 50)
            api_args.append("📋 题目文件（来自题目上传区域）：")
            api_args.append("=" * 50)
            for file_path in question_files:
                api_args.append(f"\n文件：{os.path.basename(file_path)}")
                content_type, content = process_file_content(file_path)
                if content_type == 'error':
                    raise ValueError(f"处理题目文件失败: {content}")
                elif content_type in ['image', 'pdf']:
                    api_args.append(file_path)
                else:
                    api_args.append(content)
        
        if answer_files:
            api_args.append("\n" + "=" * 50)
            api_args.append("✏️ 学生答案（来自学生答案上传区域）：")
            api_args.append("=" * 50)
            for file_path in answer_files:
                api_args.append(f"\n文件：{os.path.basename(file_path)}")
                content_type, content = process_file_content(file_path)
                if content_type == 'error':
                    raise ValueError(f"处理学生答案文件失败: {content}")
                elif content_type in ['image', 'pdf']:
                    api_args.append(file_path)
                else:
                    api_args.append(content)
        
        if marking_files:
            api_args.append("\n" + "=" * 50)
            api_args.append("📊 批改方案（来自批改方案上传区域）：")
            api_args.append("=" * 50)
            for file_path in marking_files:
                api_args.append(f"\n文件：{os.path.basename(file_path)}")
                content_type, content = process_file_content(file_path)
                if content_type == 'error':
                    raise ValueError(f"处理批改方案文件失败: {content}")
                elif content_type in ['image', 'pdf']:
                    api_args.append(file_path)
                else:
                    api_args.append(content)
        
        api_args.append("\n" + "=" * 50)
        api_args.append("\n【批改指令】：")
        api_args.append(prompt)
        
        result = api(*api_args, system_message=system_msg)
        
        return {
            "success": True,
            "result": result,
            "mode": "simplified_fixed",
            "strictness": strictness_level,
            "file_classification": "automatic_by_filename"
        }
        
    except Exception as e:
        logger.error(f"简化批改失败: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "message": "批改过程中发生错误"
        }

# ===================== 向后兼容函数 =====================
def correction_single_group(*image_files, strictness_level="中等", api=default_api, group_index=1):
    """对单个文件组进行批改"""
    try:
        prompt = prompts_module.get_complete_grading_prompt(file_info_list=[])
        return api(prompt, *image_files, system_message=prompts_module.get_core_grading_prompt())
    except Exception as e:
        error_msg = f"第{group_index}题批改失败"
        raise RuntimeError(f"{error_msg}: {str(e)}") from e

def efficient_correction_single(*image_files, strictness_level="中等", api=default_api):
    """高效简洁批改函数"""
    try:
        detailed_result = correction_single_group(*image_files, strictness_level=strictness_level, api=api)
        if len(detailed_result) > 500:
            prompt = f"""请将以下详细批改结果简化为高效简洁的格式，保留关键信息：
{detailed_result}
要求：
1. 保留题目编号和得分
2. 简要说明主要错误
3. 给出改进建议"""
            simplified = api(prompt, system_message=prompts_module.get_core_grading_prompt())
            return simplified
        return detailed_result
    except Exception as e:
        error_msg = "高效批改失败"
        raise RuntimeError(f"{error_msg}: {str(e)}") from e

def batch_efficient_correction(*image_files, strictness_level="中等", api=default_api):
    """批量高效批改函数"""
    try:
        from datetime import datetime
        
        results = []
        total_files = len(image_files)
        
        for i, file in enumerate(image_files, 1):
            try:
                result = efficient_correction_single(file, 
                                                   strictness_level=strictness_level, 
                                                   api=api)
                
                file_name = getattr(file, 'name', f'文件{i}')
                header = f"## 📄 {file_name} ({i}/{total_files})\n\n"
                results.append(header + result)
                
            except Exception as e:
                error_msg = f"文件 {i} 批改失败: {str(e)}"
                results.append(f"## ❌ 文件 {i}\n{error_msg}")
        
        final_result = "\n\n---\n\n".join(results)
        summary_header = f"\n\n# 📊 批改总览\n**共批改 {total_files} 份作业**\n✅ 批改完成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        return final_result + summary_header
        
    except Exception as e:
        error_msg = "批量批改失败"
        raise RuntimeError(f"{error_msg}: {str(e)}") from e

def generate_marking_scheme(*image_file, api=default_api):
    """生成评分方案"""
    try:
        prompt = "请根据题目生成详细的评分标准，包括每个步骤的分值分配。"
        result = api(prompt, *image_file, system_message=prompts_module.get_core_grading_prompt())
        return result
    except Exception as e:
        error_msg = "生成评分方案失败"
        raise RuntimeError(f"{error_msg}: {str(e)}") from e

def correction_with_marking_scheme(marking_schemes, student_answers, strictness_level="中等", api=default_api):
    """使用评分方案进行批改（向后兼容）"""
    try:
        if isinstance(marking_schemes, (tuple, list)):
            marking_schemes = list(marking_schemes)
        else:
            marking_schemes = [marking_schemes]
        
        if isinstance(student_answers, (tuple, list)):
            student_answers = list(student_answers)
        else:
            student_answers = [student_answers]
        
        result = batch_correction_with_standard(
            marking_schemes,
            student_answers,
            strictness_level=strictness_level,
            api=api
        )
        return result.get('correction_result', result)
    except Exception as e:
        error_msg = "批改失败"
        raise RuntimeError(f"{error_msg}: {str(e)}") from e

def correction_without_marking_scheme(student_answer, strictness_level="中等", api=default_api):
    """不使用评分方案进行批改（向后兼容）"""
    try:
        if isinstance(student_answer, (tuple, list)):
            student_answer = list(student_answer)
        else:
            student_answer = [student_answer]
        
        result = batch_correction_without_standard(
            student_answer,
            student_answer,
            strictness_level=strictness_level,
            api=api
        )
        return result.get('correction_result', result)
    except Exception as e:
        error_msg = "批改失败"
        raise RuntimeError(f"{error_msg}: {str(e)}") from e

# ===================== Web接口函数 =====================
@safe_api_call
def web_correction_with_scheme(marking_scheme_files: List[str], student_answer_files: List[str], 
                              strictness_level: str = "中等") -> Dict[str, Any]:
    """网站版本：使用评分方案进行批改"""
    result = batch_correction_with_standard(
        marking_scheme_files, 
        student_answer_files, 
        strictness_level=strictness_level, 
        api=default_api
    )
    return {
        "grading_result": result,
        "strictness_level": strictness_level,
        "scheme_files": len(marking_scheme_files),
        "answer_files": len(student_answer_files)
    }

@safe_api_call
def web_correction_without_scheme(student_answer_files: List[str], 
                                 strictness_level: str = "中等") -> Dict[str, Any]:
    """网站版本：不使用评分方案进行批改"""
    result = batch_correction_without_standard(
        student_answer_files,
        student_answer_files,
        strictness_level=strictness_level, 
        api=default_api
    )
    return {
        "grading_result": result,
        "strictness_level": strictness_level,
        "answer_files": len(student_answer_files)
    }

@safe_api_call
def web_batch_correction(batch_requests: List[Dict[str, Any]]) -> Dict[str, Any]:
    """网站版本：批量批改处理"""
    results = []
    total_requests = len(batch_requests)
    
    for i, request in enumerate(batch_requests):
        logger.info(f"处理批量请求 {i + 1}/{total_requests}")
        
        try:
            if request.get("type") == "with_scheme":
                result = web_correction_with_scheme(
                    request["marking_scheme_files"],
                    request["student_answer_files"],
                    request.get("strictness_level", "中等")
                )
            elif request.get("type") == "without_scheme":
                result = web_correction_without_scheme(
                    request["student_answer_files"],
                    request.get("strictness_level", "中等")
                )
            else:
                result = GradingResult(success=False, error_message=f"未知的请求类型: {request.get('type')}")
            
            results.append({
                "request_index": i,
                "request_id": request.get("id", f"batch_{i}"),
                "result": result.to_dict()
            })
            
        except Exception as e:
            logger.error(f"批量请求 {i + 1} 处理失败: {str(e)}")
            results.append({
                "request_index": i,
                "request_id": request.get("id", f"batch_{i}"),
                "result": GradingResult(success=False, error_message=str(e)).to_dict()
            })
    
    success_count = sum(1 for r in results if r["result"]["success"])
    
    return {
        "total_requests": total_requests,
        "success_count": success_count,
        "failure_count": total_requests - success_count,
        "results": results
    }

# ===================== API配置函数 =====================
def get_api_status() -> Dict[str, Any]:
    """获取API状态信息"""
    return {
        "api_config": api_config.get_status(),
        "status": "ready",
        "timestamp": time.time()
    }

def update_api_config(new_config: Dict[str, Any]) -> Dict[str, Any]:
    """更新API配置"""
    try:
        if "api_key" in new_config:
            api_config.api_key = new_config["api_key"]
        if "base_url" in new_config:
            api_config.base_url = new_config["base_url"]
        if "model" in new_config:
            api_config.model = new_config["model"]
        if "max_tokens" in new_config:
            api_config.max_tokens = new_config["max_tokens"]
        if "temperature" in new_config:
            api_config.temperature = new_config["temperature"]
        if "max_retries" in new_config:
            api_config.max_retries = new_config["max_retries"]
        if "retry_delay" in new_config:
            api_config.retry_delay = new_config["retry_delay"]
        
        logger.info("API配置更新成功")
        return {"success": True, "message": "配置更新成功", "updated_config": new_config}
    except Exception as e:
        error_msg = f"配置更新失败: {str(e)}"
        logger.error(error_msg)
        return {"success": False, "error": error_msg}

# ===================== 增强版批改函数 =====================
def detect_loop_and_cleanup(result):
    """智能重复内容检测和清理 - 重新设计版"""
    if not result:
        return result
    
    lines = result.split('\n')
    cleaned_lines = []
    seen_patterns = set()  # 用于检测重复模式
    recent_lines = []  # 最近的10行，用于检测立即重复
    
    # 计算行的简化特征（用于重复检测）
    def get_line_signature(line):
        # 移除变化的部分（数字、符号），只保留结构
        import re
        simplified = re.sub(r'[0-9\+\-\*\/\=\(\)\[\]]+', 'X', line.strip())
        simplified = re.sub(r'[αβγδεζηθιπ×÷±≤≥≠≈]+', 'Y', simplified)
        return simplified
    
    for line in lines:
        line = line.strip()
        if not line:
            cleaned_lines.append("")
            continue
            
        # 获取行的特征签名
        signature = get_line_signature(line)
        
        # 检测立即重复（连续相同行）
        if line in recent_lines[-5:]:  # 检查最近5行
            print(f"⚠️ 检测到立即重复，跳过: {line[:50]}...")
            continue
            
        # 检测模式重复（相同结构的行出现过多次）
        if signature in seen_patterns and len(signature) > 10:  # 忽略太短的模式
            pattern_count = sum(1 for existing_sig in seen_patterns if existing_sig == signature)
            if pattern_count > 3:  # 同样的模式超过3次
                print(f"⚠️ 检测到模式重复，跳过: {signature[:30]}...")
                continue
                
        # 检测大段重复（整个计算过程重复）
        if ('的体积' in line or '的半径' in line) and any(similar_content in prev for prev in cleaned_lines[-20:] for similar_content in ['的体积', '的半径']):
            # 计算最近20行中数学计算的密度
            math_density = sum(1 for prev in cleaned_lines[-20:] if ('的体积' in prev or '的半径' in prev or '=' in prev))
            if math_density > 15:  # 如果数学计算过于密集
                print("⚠️ 检测到大段计算重复，终止处理")
                break
            
        # 检测数学公式的过度重复
        if ('体积' in line or '半径' in line) and line.count('=') > 0:
            similar_count = sum(1 for prev in cleaned_lines[-50:] 
                              if ('体积' in prev or '半径' in prev) and prev.count('=') > 0)
            if similar_count > 10:  # 数学公式行超过10行
                print("⚠️ 检测到数学公式过度重复，停止添加")
                break
        
        # 记录模式和最近行
        seen_patterns.add(signature)
        recent_lines.append(line)
        if len(recent_lines) > 10:
            recent_lines.pop(0)
            
            cleaned_lines.append(line)
    
        # 安全截断：防止输出过长
        if len(cleaned_lines) > 200:
            cleaned_lines.append("🛑 内容过长，已自动截断")
            break
    
    print(f"✅ 循环清理完成: {len(lines)} → {len(cleaned_lines)} 行")
    return '\n'.join(cleaned_lines)

def check_question_exists(content, question_number):
    """检查题目是否存在于学生答案中"""
    if not content:
        return False
    
    content_str = str(content).strip()
    if not content_str:
        return False
    
    import re
    
    # 检查多种题目标识格式，使用正则表达式精确匹配
    patterns = [
        rf"题目\s*{question_number}(?:\s|：|:|$)",
        rf"第\s*{question_number}\s*题",
        rf"Q\s*{question_number}(?:\s|\.|\)|:|$)",
        rf"Q\.\s*{question_number}(?:\s|\.|\)|:|$)", 
        rf"问题\s*{question_number}(?:\s|：|:|$)",
        rf"^\s*{question_number}\s*[\.\)）]",  # 开头是数字加点或括号
        rf"\(\s*{question_number}\s*\)",     # (1)格式
        rf"\[\s*{question_number}\s*\]"      # [1]格式
    ]
    
    # 检查是否匹配任一模式
    for pattern in patterns:
        if re.search(pattern, content_str, re.IGNORECASE):
            return True
    
    return False

def generate_summary_with_api(grading_results):
    """使用额外API调用生成总结"""
    try:
        from .prompts_simplified import get_summary_generation_prompt
        
        # 分析批改结果
        total_questions = 0
        answered_questions = 0
        total_score = 0
        total_possible = 0
        
        # 统计信息
        for result in grading_results:
            if not result:
                continue
                
            lines = result.split('\n')
            for line in lines:
                if line.startswith('### 题目'):
                    total_questions += 1
                elif '**状态**：未作答' in line:
                    # 未作答题目不计入分数
                    continue
                elif '**得分**：' in line and '**满分**：' in lines[lines.index(line)-1]:
                    # 提取得分信息
                    try:
                        score_line = line
                        full_line = lines[lines.index(line)-1]
                        
                        # 解析得分
                        if '/' in score_line:
                            parts = score_line.split('/')
                            if len(parts) >= 2:
                                current_score = float(parts[0].split('：')[1].strip())
                                total_score += current_score
                                answered_questions += 1
                        
                        # 解析满分
                        if '：' in full_line:
                            full_score = full_line.split('：')[1].split('分')[0].strip()
                            total_possible += float(full_score)
                    except:
                        pass
        
        # 构建总结提示
        summary_prompt = f"""{get_summary_generation_prompt()}

【批改数据】
- 总题目数：{total_questions}
- 已作答题目：{answered_questions}
- 未作答题目：{total_questions - answered_questions}
- 实际得分：{total_score}
- 总可能分数：{total_possible}

【批改结果】
{chr(10).join(grading_results)}

请基于以上数据生成简洁总结："""
        
        # 调用API生成总结（不使用分批）
        logger.info("📊 开始生成批改总结...")
        logger.info(f"批改统计数据: 总题目{total_questions}, 已答{answered_questions}, 得分{total_score}/{total_possible}")
        
        summary = call_tongyiqianwen_api(summary_prompt, "")
        
        if summary:
            logger.info("📊 总结生成成功")
            logger.info(f"总结内容长度: {len(summary)} 字符")
        else:
            logger.warning("📊 总结生成失败")
            
        return summary if summary else "📊 总结生成失败"
        
    except Exception as e:
        print(f"生成总结失败: {e}")
        return "📊 总结生成出错"

def analyze_questions(content_list, file_info_list=None):
    """分析题目数量和分值 - 增强版"""
    try:
        # 直接使用题目分析提示词，避免导入问题
        analysis_prompt = """📊 题目分析任务

请仔细分析提供的文件，识别：
1. 总共有多少道题目（包括所有子题，如 (a), (b), (c) 等）
2. 每道题的分值
3. 总分是多少

【输出格式要求】
题目总数：X题
题目列表：
- 题目1：Y分
- 题目2：Z分
...
总分：W分

请严格按照上述格式输出，不要进行批改。"""
        
        # 读取所有文件内容，但保留文件路径信息
        all_contents = []
        file_paths = []
        text_contents = []
        
        for file_path in content_list:
            try:
                content_type, content = process_file_content(file_path)
                if content:
                    # 对于图片和PDF，保存文件路径
                    if content_type in ['image', 'pdf']:
                        all_contents.append(f"[文件: {os.path.basename(file_path)}]")
                        file_paths.append(file_path)
                    else:
                        all_contents.append(content)
                        file_paths.append(None)
                        text_contents.append(content)
            except Exception as e:
                print(f"⚠️ 读取文件出错 {file_path}: {e}")
        
        if not all_contents:
            print("❌ 没有可读取的文件内容")
            return None
        
        # 首先尝试从文本内容中直接识别题目（快速识别）
        total_questions = 0
        total_score = 0
        
        for content in text_contents:
            if content:
                # 识别题目模式
                import re
                # 识别题目编号模式: "题目1", "1.", "题1", "第1题", "(1)", "Question 1" 等
                question_patterns = [
                    r'题目\s*(\d+)',  # 题目1, 题目2
                    r'第\s*(\d+)\s*题',  # 第1题, 第2题
                    r'^\s*(\d+)[\.\)]\s*',  # 1. 2. 或 1) 2)
                    r'题\s*(\d+)',  # 题1, 题2
                    r'\(\s*(\d+)\s*\)',  # (1), (2)
                    r'Question\s+(\d+)',  # Question 1, Question 2
                ]
                
                question_numbers = set()
                for pattern in question_patterns:
                    matches = re.findall(pattern, content, re.MULTILINE | re.IGNORECASE)
                    for match in matches:
                        try:
                            question_numbers.add(int(match))
                        except:
                            continue
                
                if question_numbers:
                    max_question = max(question_numbers)
                    total_questions = max(total_questions, max_question)
                
                # 识别分值模式
                score_patterns = [
                    r'(\d+)\s*分',  # 3分, 5分
                    r'(\d+)\s*marks?',  # 3 marks, 5 mark
                    r'(\d+)\s*pts?',  # 3 pts, 5 pt
                    r'\[\s*(\d+)\s*M',  # [3M], [5M]
                    r'满分[：:]\s*(\d+)',  # 满分：10
                    r'总分[：:]\s*(\d+)',  # 总分：20
                ]
                
                scores = []
                for pattern in score_patterns:
                    matches = re.findall(pattern, content, re.IGNORECASE)
                    for match in matches:
                        try:
                            scores.append(int(match))
                        except:
                            continue
                
                if scores:
                    # 如果找到总分标记，使用最大的；否则累加所有分值
                    if any('总分' in content for pattern in score_patterns if re.search(pattern, content)):
                        total_score = max(total_score, max(scores))
                    else:
                        total_score = max(total_score, sum(scores))
        
        # 如果快速识别成功，直接返回结果
        if total_questions > 0:
            print(f"📊 快速识别成功：共{total_questions}题，总分{total_score}分")
            return {
                'total_questions': total_questions,
                'total_score': total_score,
                'analysis': f"快速识别：题目总数：{total_questions}题，总分：{total_score}分"
            }
        
        # 如果快速识别失败，调用API进行智能分析
        print("🤖 使用AI进行题目分析...")
        
        # 构建API参数
        api_args = [analysis_prompt]
        for i, file_path in enumerate(content_list):
            if file_paths[i]:  # 如果是图片或PDF文件
                api_args.append(file_paths[i])
            else:  # 如果是文本内容
                api_args.append(all_contents[i])
        
        result = call_tongyiqianwen_api(*api_args)
        
        if result and "API配置错误" not in result and "认证失败" not in result:
            print(f"API分析结果: {result[:200]}...")
            
            # 解析结果 - 使用更灵活的正则表达式
            import re
            
            # 提取题目总数 - 支持多种格式
            total_patterns = [
                r'题目总数[：:]\s*(\d+)',
                r'总共[：:]\s*(\d+)\s*题',
                r'共[：:]\s*(\d+)\s*题',
                r'(\d+)\s*题',
                r'总题数[：:]\s*(\d+)',
            ]
            
            total_questions = 0
            for pattern in total_patterns:
                match = re.search(pattern, result)
                if match:
                    total_questions = int(match.group(1))
                    break
            
            # 提取总分 - 支持多种格式
            score_patterns = [
                r'总分[：:]\s*(\d+)',
                r'总共[：:]\s*(\d+)\s*分',
                r'共[：:]\s*(\d+)\s*分',
                r'(\d+)\s*分\s*$',
            ]
            
            total_score = 0
            for pattern in score_patterns:
                match = re.search(pattern, result)
                if match:
                    total_score = int(match.group(1))
                    break
            
            # 如果还是没有识别到，尝试从题目列表中计算
            if total_questions == 0 or total_score == 0:
                question_list_matches = re.findall(r'-\s*题目\d+[：:]\s*(\d+)\s*分', result)
                if question_list_matches:
                    if total_questions == 0:
                        total_questions = len(question_list_matches)
                    if total_score == 0:
                        total_score = sum(int(score) for score in question_list_matches)
            
            print(f"📊 AI识别结果：共{total_questions}题，总分{total_score}分")
            return {
                'total_questions': total_questions,
                'total_score': total_score,
                'analysis': result
            }
        else:
            print("⚠️ API调用失败，使用默认值")
            return {
                'total_questions': 1,  # 默认至少有1题
                'total_score': 10,     # 默认总分10分
                'analysis': "API调用失败，使用默认设置"
            }
        
    except Exception as e:
        print(f"题目分析出错: {e}")
        import traceback
        traceback.print_exc()
        # 返回默认值而不是None，确保批改流程能继续
        return {
            'total_questions': 1,
            'total_score': 10,
            'analysis': f"分析出错: {e}"
        }

def quick_analyze_files(content_list, file_paths, file_info_list=None):
    """
    轻量级文件分析 - 第一次API调用
    只识别题目数量、学生信息等基本信息，不输出详细内容
    """
    try:
        logger.info("🔍 开始轻量级文件分析...")
        
        # 构建简洁的分析提示词
        analysis_prompt = """你是一个文件分析助手。请分析提供的文件，但只输出以下信息：

【重要】：你的任务是分析识别，不是批改作业，所以不要输出任何题目的详细内容或答案。

请按以下格式输出：

## 📊 文件分析结果

### 题目数量
总题目数：[数字]

### 学生信息
学生姓名/ID：[从文件名或内容中识别，如果有多个学生请列出]
班级信息：[从文件名、目录名或内容中识别班级信息]
学生数量：[如果是多个学生的作业]

### 文件类型
- 题目文件：[数量]个
- 答案文件：[数量]个  
- 批改标准：[数量]个

### 分析总结
[简短总结：这批文件包含什么内容，预计批改时间等]

【严格要求】：
1. 不要输出任何题目的具体内容
2. 不要输出任何答案或解题过程
3. 只输出文件的基本信息和结构分析
4. 输出要简洁明了，不超过200字"""

        # 准备API参数
        api_args = [analysis_prompt]
        
        # 添加文件（优先使用前几个文件进行快速分析）
        files_to_analyze = min(3, len(content_list))  # 最多分析前3个文件
        for i in range(files_to_analyze):
            if i < len(file_paths) and file_paths[i]:
                api_args.append(file_paths[i])
                logger.info(f"添加分析文件: {os.path.basename(file_paths[i])}")
        
        logger.info(f"执行轻量级分析 - 分析文件数: {files_to_analyze}")
        analysis_result = call_tongyiqianwen_api(*api_args)
        
        if analysis_result:
            logger.info("✅ 文件分析完成")
            logger.info(f"分析结果长度: {len(analysis_result)} 字符")
            
            # 提取关键信息
            import re
            total_questions_match = re.search(r'总题目数[：:]\s*(\d+)', analysis_result)
            student_info_match = re.search(r'学生姓名/ID[：:]\s*([^\n]+)', analysis_result)
            class_info_match = re.search(r'班级信息[：:]\s*([^\n]+)', analysis_result)
            student_count_match = re.search(r'学生数量[：:]\s*([^\n]+)', analysis_result)
            
            extracted_info = {
                'total_questions': int(total_questions_match.group(1)) if total_questions_match else 0,
                'student_info': student_info_match.group(1).strip() if student_info_match else "未识别",
                'class_info': class_info_match.group(1).strip() if class_info_match else "未识别",
                'student_count': student_count_match.group(1).strip() if student_count_match else "1",
                'analysis_text': analysis_result,
                'files_analyzed': files_to_analyze
            }
            
            logger.info(f"提取信息: 题目数={extracted_info['total_questions']}, 学生={extracted_info['student_info']}, 班级={extracted_info['class_info']}")
            return extracted_info
        else:
            logger.warning("⚠️ 文件分析失败，使用默认信息")
            return {
                'total_questions': len(content_list),
                'student_info': "未识别",
                'class_info': "未识别",
                'student_count': "1",
                'analysis_text': "分析失败",
                'files_analyzed': 0
            }
            
    except Exception as e:
        logger.error(f"❌ 文件分析出错: {e}")
        return {
            'total_questions': len(content_list),
            'student_info': "分析出错",
            'class_info': "分析出错",
            'student_count': "1",
            'analysis_text': f"分析出错: {e}",
            'files_analyzed': 0
        }

def enhanced_batch_correction_with_standard(content_list, file_info_list=None, batch_size=10, generate_summary=True):
    """增强版分批批改（有标准答案）- 超级严格版"""
    try:
        logger.info("🚀 开始增强版批改（有标准答案）")
        logger.info(f"输入文件数量: {len(content_list)}")
        logger.info(f"生成总结: {generate_summary}")
        
        from .prompts_simplified import get_core_grading_prompt, get_batch_processing_prompt
        
        # 首先读取所有文件内容，但保留文件路径信息
        all_contents = []
        file_paths = []  # 保存原始文件路径
        for i, file_path in enumerate(content_list):
            try:
                logger.info(f"处理文件 {i+1}/{len(content_list)}: {os.path.basename(file_path)}")
                content_type, content = process_file_content(file_path)
                
                # 检查是否处理成功
                if content_type == 'error':
                    logger.warning(f"⚠️ 文件处理失败：{file_path} - {content}")
                    continue
                
                if content and content_type != 'error':
                    # 对于图片和PDF，保存文件路径而不是内容
                    if content_type in ['image', 'pdf']:
                        all_contents.append(f"[文件: {os.path.basename(file_path)}]")
                        file_paths.append(file_path)
                        logger.info(f"文件处理成功: {os.path.basename(file_path)} (类型: {content_type})")
                    else:
                        all_contents.append(content)
                        file_paths.append(None)  # 文本内容不需要文件路径
                        logger.info(f"文件处理成功: {os.path.basename(file_path)} (类型: {content_type}, 长度: {len(content)} 字符)")
                else:
                    logger.warning(f"⚠️ 无法读取文件：{file_path}")
            except Exception as e:
                logger.error(f"⚠️ 读取文件出错 {file_path}: {e}")
        
        if not all_contents:
            logger.error("❌ 没有可处理的文件内容")
            logger.error(f"调试信息：")
            logger.error(f"- 输入文件数量: {len(content_list)}")
            logger.error(f"- 文件列表: {[os.path.basename(f) if isinstance(f, str) and os.path.exists(f) else str(f) for f in content_list]}")
            logger.error(f"- 处理结果列表长度: {len(all_contents)}")
            logger.error(f"- 文件路径列表长度: {len(file_paths)}")
            
            # 提供更详细的错误信息
            return {
                'text': "❌ 文件处理失败：没有可处理的文件内容。请检查：\n1. 文件是否正确上传\n2. 文件格式是否支持\n3. 文件是否损坏",
                'html': "<div style='color: red;'>❌ 文件处理失败：没有可处理的文件内容。请检查：<br>1. 文件是否正确上传<br>2. 文件格式是否支持<br>3. 文件是否损坏</div>",
                'format': 'error'
            }
        
        logger.info(f"✅ 文件处理完成，共处理 {len(all_contents)} 个文件")
        
        # 合并所有内容（仅用于显示）
        combined_contents = []
        for content in all_contents:
            if isinstance(content, tuple) and len(content) == 2:
                # 如果是元组，取第二个元素（内容）
                combined_contents.append(str(content[1]))
            else:
                combined_contents.append(str(content))
        combined_content = "\n\n".join(combined_contents)
        
        # 第一步：识别题目数量和学生信息（轻量级API调用）
        logger.info("🔍 第一步：识别题目数量和学生信息...")
        analysis_result = quick_analyze_files(content_list, file_paths, file_info_list)
        
        logger.info("📝 第二步：开始构建批改提示词...")
        prompt = get_core_grading_prompt(file_info_list, analysis_result)
        
        if file_info_list:
            marking_content = extract_marking_content(file_info_list)
            if marking_content:
                prompt = f"{prompt}\n\n批改标准：\n{marking_content}"
                logger.info(f"已添加批改标准，长度: {len(marking_content)} 字符")
        
        # 调用API - 传递文件路径而不是文本内容
        api_args = [prompt]
        for i, file_path in enumerate(content_list):
            if file_paths[i]:  # 如果是图片或PDF文件
                api_args.append(file_paths[i])
                logger.info(f"添加API参数: 文件路径 {os.path.basename(file_paths[i])}")
            else:  # 如果是文本内容
                api_args.append(all_contents[i])
                logger.info(f"添加API参数: 文本内容 (长度: {len(all_contents[i])} 字符)")
        
        logger.info(f"调用API - 总参数数量: {len(api_args)}")
        result = call_tongyiqianwen_api(*api_args)
        
        if result:
            logger.info("🎯 API调用成功，开始结果处理流程")
            logger.info(f"原始结果长度: {len(result)} 字符")
            
            # 1. 强制数学符号转换（完全重塑版）
            logger.info("1. 执行数学符号转换...")
            result = convert_latex_to_unicode(result)
            
            # 2. 终极循环检测和清理
            logger.info("2. 执行循环检测和清理...")
            result = detect_loop_and_cleanup(result)
            
            # 3. 移除任何总结内容
            logger.info("3. 移除总结内容...")
            result = remove_summary_from_batch(result)
            
            # 4. 强制格式修正
            logger.info("4. 执行格式修正...")
            result = enforce_strict_format(result)
            
            # 5. 清理混乱输出（移除重复题目和思考内容）
            logger.info("5. 清理混乱输出...")
            result = clean_grading_output(result)
            
            # 6. 强制分值限制（防止超出题目总分）
            logger.info("6. 执行分值限制检查...")
            if file_info_list:
                marking_content = extract_marking_content(file_info_list)
                if marking_content:
                    result = enforce_score_limits(result, marking_content)
                    logger.info("分值限制检查完成")
            
            logger.info(f"结果处理完成，最终长度: {len(result)} 字符")
            
            # 如果需要总结，进行额外的API调用
            if generate_summary:
                logger.info("📊 开始生成批改总结...")
                summary = generate_summary_with_api([result])
                result += f"\n\n{summary}"
                logger.info("总结生成完成")
            
            # 转换为HTML格式
            logger.info("🎨 开始转换为HTML格式...")
            html_result = convert_to_html_markdown(result)
            logger.info(f"HTML转换完成，长度: {len(html_result)} 字符")
            
            # 返回包含原始结果和HTML结果的字典
            logger.info("✅ 批改流程完全完成")
            return {
                'text': result,
                'html': html_result,
                'format': 'enhanced'
            }
        else:
            logger.error("❌ API调用返回空结果")
            return None

        
    except Exception as e:
        logger.error(f"❌ 增强版批改（有标准）出错: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return None

def enhanced_batch_correction_without_standard(content_list, file_info_list=None, batch_size=10, generate_summary=True):
    """增强版分批批改（无标准答案）- 超级严格版"""
    try:
        from .prompts_simplified import get_core_grading_prompt, get_batch_processing_prompt
        
        # 首先读取所有文件内容
        all_contents = []
        file_paths = []  # 保存原始文件路径
        for file_path in content_list:
            try:
                content_type, content = process_file_content(file_path)
                
                # 检查是否处理成功
                if content_type == 'error':
                    print(f"⚠️ 文件处理失败：{file_path} - {content}")
                    continue
                
                if content and content_type != 'error':
                    # 对于图片和PDF，保存文件路径而不是内容
                    if content_type in ['image', 'pdf']:
                        all_contents.append(f"[文件: {os.path.basename(file_path)}]")
                        file_paths.append(file_path)
                    else:
                        all_contents.append(content)
                        file_paths.append(None)
                else:
                    print(f"⚠️ 无法读取文件：{file_path}")
            except Exception as e:
                print(f"⚠️ 读取文件出错 {file_path}: {e}")
        
        if not all_contents:
            error_msg = "❌ 没有可处理的文件内容"
            print(error_msg)
            logger.error(error_msg)
            logger.error(f"调试信息：")
            logger.error(f"- 输入文件数量: {len(content_list)}")
            logger.error(f"- 文件列表: {[os.path.basename(f) if isinstance(f, str) and os.path.exists(f) else str(f) for f in content_list]}")
            logger.error(f"- 处理结果列表长度: {len(all_contents)}")
            logger.error(f"- 文件路径列表长度: {len(file_paths)}")
            
            # 提供更详细的错误信息
            return {
                'text': "❌ 文件处理失败：没有可处理的文件内容。请检查：\n1. 文件是否正确上传\n2. 文件格式是否支持\n3. 文件是否损坏",
                'html': "<div style='color: red;'>❌ 文件处理失败：没有可处理的文件内容。请检查：<br>1. 文件是否正确上传<br>2. 文件格式是否支持<br>3. 文件是否损坏</div>",
                'format': 'error'
            }
        
        # 合并所有内容
        combined_contents = []
        for content in all_contents:
            if isinstance(content, tuple) and len(content) == 2:
                # 如果是元组，取第二个元素（内容）
                combined_contents.append(str(content[1]))
            else:
                combined_contents.append(str(content))
        combined_content = "\n\n".join(combined_contents)
        
        # 第一步：识别题目数量和学生信息（轻量级API调用）
        print("🔍 第一步：识别题目数量和学生信息...")
        analysis_result = quick_analyze_files(content_list, file_paths, file_info_list)
        
        # 第二步：调用API进行批改
        print("📝 第二步：开始构建批改提示词...")
        prompt = get_core_grading_prompt(file_info_list, analysis_result)
        
        # 调用API - 传递文件路径而不是文本内容
        api_args = [prompt]
        for i, file_path in enumerate(content_list):
            if i < len(file_paths) and file_paths[i]:  # 如果是图片或PDF文件
                api_args.append(file_paths[i])
                print(f"添加API参数: 文件路径 {os.path.basename(file_paths[i])}")
            elif i < len(all_contents):  # 如果是文本内容
                api_args.append(all_contents[i])
                print(f"添加API参数: 文本内容 (长度: {len(all_contents[i])} 字符)")
        
        print(f"调用API - 总参数数量: {len(api_args)}")
        result = call_tongyiqianwen_api(*api_args)
        
        if result:
            # 1. 强制数学符号转换（完全重塑版）
            result = convert_latex_to_unicode(result)
            
            # 2. 终极循环检测和清理
            result = detect_loop_and_cleanup(result)
            
            # 3. 移除任何总结内容
            result = remove_summary_from_batch(result)
            
            # 4. 强制格式修正
            result = enforce_strict_format(result)
            
            # 5. 清理混乱输出（移除重复题目和思考内容）
            result = clean_grading_output(result)
            
            # 如果需要总结，进行额外的API调用
            if generate_summary:
                print("📊 正在生成批改总结...")
                summary = generate_summary_with_api([result])
                result += f"\n\n{summary}"
            
            # 转换为HTML格式
            print("🎨 正在转换为HTML格式...")
            html_result = convert_to_html_markdown(result)
            
            # 返回包含原始结果和HTML结果的字典
            return {
                'text': result,
                'html': html_result,
                'format': 'enhanced'
            }

        
    except Exception as e:
        logger.error(f"❌ 增强版批改（无标准）出错: {e}")
        import traceback
        logger.error(f"错误详情: {traceback.format_exc()}")
        return None

def remove_summary_from_batch(result):
    """移除批次结果中的总结内容"""
    if not result:
        return result
    
    lines = result.split('\n')
    cleaned_lines = []
    
    # 检测总结关键词
    summary_keywords = [
        '总结', '汇总', '总体', '整体表现', '综合评价', 
        '批改完成', '总得分', '平均分', '建议'
    ]
    
    skip_mode = False
    for line in lines:
        line_lower = line.lower()
        
        # 检测是否进入总结模式
        if any(keyword in line for keyword in summary_keywords):
            skip_mode = True
            continue
            
        # 如果不在总结模式，保留该行
        if not skip_mode:
            cleaned_lines.append(line)
        else:
            # 检测是否退出总结模式（遇到新题目）
            if line.startswith('### 题目'):
                skip_mode = False
                cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def enforce_strict_format(result):
    """强制执行严格格式 - 完全重塑版"""
    if not result:
        return result
    
    import re
    
    # 第一步：移除所有^符号
    result = result.replace('^', '')
    
    # 第二步：预处理 - 在题目标记前后添加换行
    result = re.sub(r'([^\n])(### 题目)', r'\1\n\n\2', result)
    result = re.sub(r'(### 题目\d+)([^\n])', r'\1\n', result)
    
    # 第三步：修复标题格式
    result = re.sub(r'(\*\*)?满分(\*\*)?[：:]', '**满分**：', result)
    result = re.sub(r'(\*\*)?得分(\*\*)?[：:]', '**得分**：', result)
    result = re.sub(r'(\*\*)?批改详情(\*\*)?[：:]', '**批改详情**：', result)
    
    # 第四步：确保每个字段独占一行
    result = re.sub(r'(\d+分)\s*(- 📊[^*]+)\s*(\*\*得分)', r'\1 \2\n\3', result)
    result = re.sub(r'(\d+分)\s*(\*\*批改详情)', r'\1\n\2', result)
    result = re.sub(r'(：\d+分)\s*(-)', r'\1\n\2', result)
    
    # 修复双破折号问题
    result = re.sub(r'- - 📊', '- 📊', result)
    
    # 第五步：按题目分割处理
    questions = re.split(r'(### 题目\d+)', result)
    
    formatted_parts = []
    for i in range(1, len(questions), 2):
        if i < len(questions) - 1:
            title = questions[i]
            content = questions[i + 1]
            
            # 处理题目内容
            lines = content.split('\n')
            formatted_lines = [title]
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 处理各个字段
                if line.startswith('**满分**') or line.startswith('**得分**') or line.startswith('**批改详情**'):
                    formatted_lines.append(line)
                elif '✓' in line or '✗' in line or (line.startswith('-') and '[' in line and ']' in line):
                    # 处理批改步骤
                    # 检查是否有多个步骤在一行
                    if line.count('✓') + line.count('✗') > 1:
                        # 分割多个步骤
                        # 使用正则表达式找到所有步骤
                        step_pattern = r'(-\s*(?:\([a-zA-Z]\)\s*)?[^✓✗]+(?:✓|✗)\s*\[\d+[MA]/A\])'
                        steps = re.findall(step_pattern, line)
                        
                        if steps:
                            # 逐个处理每个步骤
                            for step in steps:
                                step = step.strip()
                                # 确保格式正确
                                if not step.startswith('- '):
                                    step = '- ' + step.lstrip('-').strip()
                                
                                # 处理 (a), (b) 等标记
                                step = re.sub(r'-\s*\(([a-zA-Z])\)\s*', r'- (\1) ', step)
                                
                                # 确保符号前有空格
                                step = re.sub(r'([^\s])(✓|✗)', r'\1 \2', step)
                                
                                # 确保分值格式正确
                                step = re.sub(r'\[\s*(\d+)\s*([MA])\s*/?\s*([MA])?\s*\]', r'[\1\2/A]', step)
                                step = re.sub(r'\[\s*(\d+)\s*([MA])\s*\]', r'[\1\2/A]', step)
                                
                                formatted_lines.append(step)
                        else:
                            # 如果正则匹配失败，尝试按符号分割
                            parts = re.split(r'(?=[✓✗])', line)
                            for i, part in enumerate(parts):
                                if part.strip() and ('✓' in part or '✗' in part):
                                    part = part.strip()
                                    if not part.startswith('- '):
                                        part = '- ' + part
                                    formatted_lines.append(part)
                    else:
                        # 单个步骤
                        line = re.sub(r'\s+', ' ', line)
                        
                        # 确保步骤格式正确
                        if not line.startswith('- '):
                            if line.startswith('-'):
                                line = '- ' + line[1:].strip()
                            else:
                                line = '- ' + line
                        
                        # 处理 (a), (b) 等标记
                        line = re.sub(r'-\s*\(([a-zA-Z])\)\s*', r'- (\1) ', line)
                        
                        # 确保符号前有空格
                        line = re.sub(r'([^\s])(✓|✗)', r'\1 \2', line)
                        
                        # 确保分值格式正确 [XM/A]
                        line = re.sub(r'\[\s*(\d+)\s*([MA])\s*/?\s*([MA])?\s*\]', r'[\1\2/A]', line)
                        line = re.sub(r'\[\s*(\d+)\s*([MA])\s*\]', r'[\1\2/A]', line)
                        
                        formatted_lines.append(line)
                elif '📊 来源：MARKING标准' in line:
                    # 这部分应该跟在满分后面
                    if formatted_lines and '**满分**' in formatted_lines[-1]:
                        # 检查是否已经有破折号
                        if formatted_lines[-1].endswith(' - '):
                            formatted_lines[-1] = formatted_lines[-1] + line.strip()
                        else:
                            formatted_lines[-1] = formatted_lines[-1] + ' - ' + line.strip()
                else:
                    # 其他内容，如果不是总结相关的就保留
                    if not any(skip in line for skip in ['总结', '批改完成', '整体表现', '统计信息', '主要错误', '改进建议']):
                        formatted_lines.append(line)
            
            # 添加格式化后的题目
            if formatted_parts:
                formatted_parts.append('')  # 题目之间的空行
            formatted_parts.extend(formatted_lines)
    
    # 第六步：组合所有部分
    result = '\n'.join(formatted_parts)
    
    # 第七步：最终清理
    # 确保没有连续的空行
    result = re.sub(r'\n{3,}', '\n\n', result)
    
    # 确保每个题目块的格式正确
    result = re.sub(r'(### 题目\d+)\n+(\*\*满分)', r'\1\n\2', result)
    result = re.sub(r'(\*\*满分\*\*：[^\n]+)\n+(\*\*得分)', r'\1\n\2', result)
    result = re.sub(r'(\*\*得分\*\*：[^\n]+)\n+(\*\*批改详情)', r'\1\n\2', result)
    result = re.sub(r'(\*\*批改详情\*\*：)\n+(- )', r'\1\n\2', result)
    
    # 移除行首行尾的空格
    lines = result.split('\n')
    result = '\n'.join(line.rstrip() for line in lines)
    
    # 确保开头没有空行
    result = result.lstrip('\n')
    
    return result

# ===================== 增强版批改函数 =====================
def convert_latex_to_unicode(text):
    """完全重塑的LaTeX到Unicode转换系统"""
    if not text:
        return text
    
    # 第一步：检测LaTeX符号
    if '$' in text:
        print("⚠️ 警告：检测到LaTeX符号，正在强制转换...")
    
    # 第二步：移除所有$符号
    text = text.replace('$', '')
    
    # 第三步：完整的LaTeX到Unicode映射表
    latex_replacements = {
        # 希腊字母
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
        r'\epsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η', r'\theta': 'θ',
        r'\iota': 'ι', r'\kappa': 'κ', r'\lambda': 'λ', r'\mu': 'μ',
        r'\nu': 'ν', r'\xi': 'ξ', r'\pi': 'π', r'\rho': 'ρ',
        r'\sigma': 'σ', r'\tau': 'τ', r'\phi': 'φ', r'\chi': 'χ',
        r'\psi': 'ψ', r'\omega': 'ω',
        r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
        r'\Xi': 'Ξ', r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Phi': 'Φ',
        r'\Psi': 'Ψ', r'\Omega': 'Ω',
        
        # 数学运算符
        r'\times': '×', r'\div': '÷', r'\pm': '±', r'\mp': '∓',
        r'\cdot': '·', r'\ast': '*', r'\star': '⋆', r'\dagger': '†',
        r'\ddagger': '‡', r'\amalg': '⨿',
        
        # 关系符号
        r'\leq': '≤', r'\geq': '≥', r'\neq': '≠', r'\approx': '≈',
        r'\equiv': '≡', r'\sim': '∼', r'\simeq': '≃', r'\cong': '≅',
        r'\propto': '∝', r'\perp': '⊥', r'\parallel': '∥',
        
        # 集合符号
        r'\in': '∈', r'\notin': '∉', r'\subset': '⊂', r'\supset': '⊃',
        r'\subseteq': '⊆', r'\supseteq': '⊇', r'\cup': '∪', r'\cap': '∩',
        r'\emptyset': '∅', r'\varnothing': '∅',
        
        # 逻辑符号
        r'\land': '∧', r'\lor': '∨', r'\lnot': '¬', r'\forall': '∀',
        r'\exists': '∃', r'\nexists': '∄', r'\therefore': '∴', r'\because': '∵',
        
        # 箭头
        r'\rightarrow': '→', r'\leftarrow': '←', r'\leftrightarrow': '↔',
        r'\Rightarrow': '⇒', r'\Leftarrow': '⇐', r'\Leftrightarrow': '⇔',
        r'\mapsto': '↦', r'\to': '→',
        
        # 其他符号
        r'\infty': '∞', r'\partial': '∂', r'\nabla': '∇', r'\square': '□',
        r'\triangle': '△', r'\angle': '∠', r'\measuredangle': '∡',
        r'\sphericalangle': '∢', r'\prime': '′', r'\circ': '°',
        r'\bullet': '•', r'\dots': '…', r'\cdots': '⋯', r'\vdots': '⋮',
        r'\ddots': '⋱', r'\ldots': '…',
        
        # 特殊函数
        r'\sin': 'sin', r'\cos': 'cos', r'\tan': 'tan', r'\cot': 'cot',
        r'\sec': 'sec', r'\csc': 'csc', r'\log': 'log', r'\ln': 'ln',
        r'\exp': 'exp', r'\min': 'min', r'\max': 'max', r'\sup': 'sup',
        r'\inf': 'inf', r'\lim': 'lim', r'\limsup': 'lim sup', r'\liminf': 'lim inf',
        
        # 分数处理（特殊处理）
        r'\frac': '/',  # 临时标记，后续特殊处理
        
        # 上下标（特殊处理）
        '^': '**', '_': '__',  # 临时标记
        
        # 括号
        r'\left': '', r'\right': '',
        r'\{': '{', r'\}': '}',
        r'\langle': '⟨', r'\rangle': '⟩',
        r'\lfloor': '⌊', r'\rfloor': '⌋',
        r'\lceil': '⌈', r'\rceil': '⌉',
        
        # 求和、积分等
        r'\sum': '∑', r'\prod': '∏', r'\int': '∫', r'\oint': '∮',
        r'\iint': '∬', r'\iiint': '∭', r'\bigcup': '⋃', r'\bigcap': '⋂',
        
        # 根号
        r'\sqrt': '√',
        
        # 矩阵
        r'\begin{pmatrix}': '(', r'\end{pmatrix}': ')',
        r'\begin{bmatrix}': '[', r'\end{bmatrix}': ']',
        r'\begin{vmatrix}': '|', r'\end{vmatrix}': '|',
        r'\begin{cases}': '{', r'\end{cases}': '',
        r'\\': '; ',  # 矩阵换行
    }
    
    # 第四步：应用所有替换
    for latex, unicode_char in latex_replacements.items():
        text = text.replace(latex, unicode_char)
    
    # 第五步：处理分数 \frac{a}{b} -> a/b
    import re
    frac_pattern = r'/\{([^}]+)\}\{([^}]+)\}'
    while re.search(frac_pattern, text):
        text = re.sub(frac_pattern, r'(\1)/(\2)', text)
    
    # 第六步：处理上标 a^{bc} -> a^(bc), a^b -> a^b
    superscript_pattern = r'\*\*\{([^}]+)\}'
    text = re.sub(superscript_pattern, r'^(\1)', text)
    text = text.replace('**', '^')
    
    # 第七步：处理下标 a_{bc} -> a_(bc), a_b -> a_b
    subscript_pattern = r'__\{([^}]+)\}'
    text = re.sub(subscript_pattern, r'_(\1)', text)
    text = text.replace('__', '_')
    
    # 第八步：清理多余的花括号
    text = re.sub(r'\{([^}]+)\}', r'\1', text)
    
    # 第九步：最终清理
    text = text.replace('\\', '')  # 移除剩余的反斜杠
    text = re.sub(r'\s+', ' ', text)  # 规范化空格
    text = text.strip()
    
    # 第十步：处理常见的数学表达式
    # 处理上标数字（转换为Unicode上标）
    superscript_map = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
        '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
        'n': 'ⁿ', 'i': 'ⁱ'
    }
    
    # 转换简单上标 x^2 -> x²
    def replace_superscript(match):
        base = match.group(1)
        sup = match.group(2)
        if len(sup) == 1 and sup in superscript_map:
            return base + superscript_map[sup]
        else:
            # 多字符上标保持原样
            return f"{base}^({sup})"
    
    text = re.sub(r'([a-zA-Z0-9])\^([0-9+\-=()ni])', replace_superscript, text)
    text = re.sub(r'([a-zA-Z0-9])\^\(([^)]+)\)', r'\1^(\2)', text)
    
    # 处理下标数字（转换为Unicode下标）
    subscript_map = {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
        '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
        'a': 'ₐ', 'e': 'ₑ', 'o': 'ₒ', 'x': 'ₓ', 'h': 'ₕ',
        'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'p': 'ₚ',
        's': 'ₛ', 't': 'ₜ'
    }
    
    # 转换简单下标 x_1 -> x₁
    def replace_subscript(match):
        base = match.group(1)
        sub = match.group(2)
        if len(sub) == 1 and sub in subscript_map:
            return base + subscript_map[sub]
        else:
            # 多字符下标保持原样
            return f"{base}_({sub})"
    
    text = re.sub(r'([a-zA-Z0-9])_([0-9+\-=()aeoxhklmnpst])', replace_subscript, text)
    text = re.sub(r'([a-zA-Z0-9])_\(([^)]+)\)', r'\1_(\2)', text)
    
    # 第十一步：验证是否还有LaTeX残留
    latex_indicators = ['\\', '$', r'\frac', r'\angle', r'\times']
    for indicator in latex_indicators:
        if indicator in text:
            print(f"⚠️ 警告：仍检测到LaTeX符号 '{indicator}'")
    
    return text

def clean_grading_output(result):
    """清理批改输出，移除混乱内容"""
    if not result:
        return result
    
    lines = result.split('\n')
    cleaned_lines = []
    seen_questions = set()
    current_question = None
    skip_mode = False
    
    for line in lines:
        # 检测批改暂停思考等混乱内容
        if '批改暂停思考' in line or '进度报告' in line or '============' in line:
            skip_mode = True
            continue
        
        # 如果在跳过模式，检查是否遇到新题目
        if skip_mode:
            if line.strip().startswith('### 题目'):
                skip_mode = False
            else:
                continue
        
        # 检测题目标题
        if line.strip().startswith('### 题目'):
            import re
            match = re.search(r'### 题目(\d+)', line)
            if match:
                question_num = int(match.group(1))
                if question_num in seen_questions:
                    # 跳过重复的题目
                    print(f"⚠️ 检测到重复题目{question_num}，跳过")
                    current_question = None
                    continue
                else:
                    seen_questions.add(question_num)
                    current_question = question_num
                    cleaned_lines.append(line)
        elif current_question is not None and not skip_mode:
            # 保留当前题目的内容
            cleaned_lines.append(line)
    
    return '\n'.join(cleaned_lines)

def enforce_score_limits(result, marking_scheme):
    """强制执行分值限制，防止超出题目总分"""
    if not result or not marking_scheme:
        return result
    
    # 解析批改标准获取每题分值
    question_scores = parse_marking_scheme(marking_scheme)
    
    lines = result.split('\n')
    corrected_lines = []
    current_question = None
    current_steps = []
    
    for line in lines:
        # 检测题目
        import re
        question_match = re.search(r'### 题目(\d+)', line)
        if question_match:
            # 处理上一题的步骤
            if current_question and current_steps:
                corrected_lines.extend(enforce_steps_limit(current_steps, question_scores.get(current_question, 0)))
            
            # 开始新题目
            current_question = int(question_match.group(1))
            current_steps = []
            corrected_lines.append(line)
        elif line.strip().startswith('关键步骤') or line.strip().startswith('- 关键步骤'):
            current_steps.append(line)
        else:
            if current_steps:
                # 完成当前题目的步骤收集
                corrected_lines.extend(enforce_steps_limit(current_steps, question_scores.get(current_question, 0)))
                current_steps = []
            corrected_lines.append(line)
    
    # 处理最后一题
    if current_steps:
        corrected_lines.extend(enforce_steps_limit(current_steps, question_scores.get(current_question, 0)))
    
    return '\n'.join(corrected_lines)

def enforce_steps_limit(steps, max_score):
    """限制步骤数量，不超过题目总分"""
    if not steps:
        return steps
    
    # 计算当前步骤的总分
    total_score = 0
    valid_steps = []
    
    for step in steps:
        # 提取步骤分值
        import re
        score_match = re.search(r'\[(\d+)[MA]\]', step)
        if score_match:
            step_score = int(score_match.group(1))
            if total_score + step_score <= max_score:
                valid_steps.append(step)
                total_score += step_score
            else:
                print(f"⚠️ 跳过超出分值的步骤：{step}")
        else:
            valid_steps.append(step)
    
    return valid_steps

def parse_marking_scheme(marking_scheme):
    """解析批改标准，获取每题分值"""
    question_scores = {}
    
    if not marking_scheme:
        return question_scores
    
    lines = marking_scheme.split('\n')
    
    import re
    for line in lines:
        # 检测题目和分值 - 支持多种格式
        patterns = [
            r'(\d+)\.\s*.*?\((\d+)\s*(?:marks?|分)\)',  # 1. xxx (3 marks)
            r'题目\s*(\d+).*?\((\d+)\s*(?:marks?|分)\)',  # 题目1 (3分)
            r'Q(\d+).*?\((\d+)\s*(?:marks?|分)\)',       # Q1 (3 marks)
            r'Question\s*(\d+).*?\((\d+)\s*(?:marks?|分)\)'  # Question 1 (3 marks)
        ]
        
        for pattern in patterns:
            match = re.search(pattern, line, re.IGNORECASE)
            if match:
                question_num = int(match.group(1))
                score = int(match.group(2))
                question_scores[question_num] = score
                break
    
    return question_scores

def extract_marking_content(file_info_list):
    """从文件信息中提取批改标准内容"""
    if not file_info_list:
        return ""
    
    marking_contents = []
    
    for file_info in file_info_list:
        if isinstance(file_info, dict):
            file_name = file_info.get('name', '').upper()
            file_path = file_info.get('path', '')
            
            # 检查是否是批改标准文件
            if 'MARKING' in file_name or 'STANDARD' in file_name or '标准' in file_name:
                try:
                    content_type, content = process_file_content(file_path)
                    if content:
                        # 如果是元组，取第二个元素
                        if isinstance(content, tuple) and len(content) == 2:
                            marking_contents.append(str(content[1]))
                        else:
                            marking_contents.append(str(content))
                except Exception as e:
                    print(f"读取批改标准文件失败 {file_path}: {e}")
    
    return "\n\n".join(marking_contents)

def convert_to_html_markdown(result):
    """将批改结果转换为HTML格式的Markdown - 增强版"""
    if not result:
        return result
    
    import re
    
    # HTML模板 - 增强样式（白色背景）
    html_template = """<div style="font-family: 'Microsoft YaHei', Arial, sans-serif; color: #333; line-height: 1.8; max-width: 800px; margin: 0 auto; padding: 20px; background-color: #ffffff;">
<style>
    body {{
        background-color: #ffffff !important;
    }}
    .question-title {{
        color: #2c3e50;
        margin-top: 25px;
        margin-bottom: 15px;
        border-bottom: 2px solid #3498db;
        padding-bottom: 8px;
        font-weight: bold;
        font-size: 1.5em;
    }}
    .score-full {{
        color: #27ae60;
        font-weight: bold;
    }}
    .score-good {{
        color: #f39c12;
        font-weight: bold;
    }}
    .score-low {{
        color: #e74c3c;
        font-weight: bold;
    }}
    .step-correct {{
        color: #27ae60;
        margin: 3px 0 3px 20px;
    }}
    .step-wrong {{
        color: #e74c3c;
        margin: 3px 0 3px 20px;
    }}
    .check-mark {{
        font-weight: bold;
        font-size: 1.2em;
    }}
    .section-title {{
        font-weight: bold;
        color: #34495e;
        margin: 10px 0 5px 0;
    }}
</style>
{content}
</div>"""
    
    # 处理每个题目
    lines = result.split('\n')
    html_lines = []
    current_full_score = 0
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            html_lines.append('<br>')
            continue
        
        # 题目标题 - 加粗
        if line.startswith('### 题目'):
            question_match = re.search(r'### 题目(\d+)', line)
            if question_match:
                q_num = question_match.group(1)
                html_lines.append(f'<h3 class="question-title">题目 {q_num}</h3>')
        
        # 满分行
        elif line.startswith('**满分**：'):
            content = line.replace('**满分**：', '').strip()
            # 提取满分数值
            score_match = re.search(r'(\d+)分', content)
            if score_match:
                current_full_score = int(score_match.group(1))
            html_lines.append(f'<p style="margin: 5px 0;"><span class="section-title">满分：</span>{content}</p>')
        
        # 得分行
        elif line.startswith('**得分**：'):
            content = line.replace('**得分**：', '').strip()
            # 根据得分情况添加不同颜色
            score_match = re.search(r'(\d+)分', content)
            if score_match and current_full_score > 0:
                score = int(score_match.group(1))
                percentage = score / current_full_score
                if percentage >= 1.0:
                    score_class = 'score-full'
                elif percentage >= 0.8:
                    score_class = 'score-good'
                else:
                    score_class = 'score-low'
                html_lines.append(f'<p style="margin: 5px 0;"><span class="section-title">得分：</span><span class="{score_class}">{content}</span></p>')
            else:
                html_lines.append(f'<p style="margin: 5px 0;"><span class="section-title">得分：</span>{content}</p>')
        
        # 批改详情标题
        elif line.startswith('**批改详情**：'):
            html_lines.append('<p class="section-title">批改详情：</p>')
        
        # 批改步骤 - 得分绿色，失分红色
        elif line.startswith('- ') and ('✓' in line or '✗' in line):
            step_content = line[2:].strip()  # 移除 "- "
            
            if '✓' in line:
                # 正确的步骤 - 绿色
                step_html = step_content.replace('✓', '<span class="check-mark" style="color: #27ae60;">✓</span>')
                # 提取步骤内容和分数
                parts = re.split(r'(✓)', step_content)
                if len(parts) >= 3:
                    before_check = parts[0].strip()
                    after_check = parts[2].strip()
                    html_lines.append(f'<p class="step-correct">• <strong>{before_check}</strong> <span class="check-mark" style="color: #27ae60;">✓</span> {after_check}</p>')
                else:
                    html_lines.append(f'<p class="step-correct">• {step_html}</p>')
            else:
                # 错误的步骤 - 红色
                step_html = step_content.replace('✗', '<span class="check-mark" style="color: #e74c3c;">✗</span>')
                # 提取步骤内容和分数
                parts = re.split(r'(✗)', step_content)
                if len(parts) >= 3:
                    before_check = parts[0].strip()
                    after_check = parts[2].strip()
                    # 检查是否有扣分原因（→ 符号）
                    if '→' in after_check:
                        score_and_reason = after_check.split('→', 1)
                        if len(score_and_reason) == 2:
                            score_part = score_and_reason[0].strip()
                            reason_part = score_and_reason[1].strip()
                            html_lines.append(f'<p class="step-wrong">• <strong>{before_check}</strong> <span class="check-mark" style="color: #e74c3c;">✗</span> {score_part} → <em>{reason_part}</em></p>')
                        else:
                            html_lines.append(f'<p class="step-wrong">• <strong>{before_check}</strong> <span class="check-mark" style="color: #e74c3c;">✗</span> {after_check}</p>')
                    else:
                        html_lines.append(f'<p class="step-wrong">• <strong>{before_check}</strong> <span class="check-mark" style="color: #e74c3c;">✗</span> {after_check}</p>')
                else:
                    html_lines.append(f'<p class="step-wrong">• {step_html}</p>')
        
        # 总结部分
        elif '=== 📊 批改总结 ===' in line:
            html_lines.append('<hr style="margin: 30px 0; border: none; border-top: 2px solid #ecf0f1;">')
            html_lines.append('<h2 style="color: #2c3e50; text-align: center; margin: 20px 0; font-weight: bold;">📊 批改总结</h2>')
        
        elif line.startswith('**整体表现**：'):
            content = line.replace('**整体表现**：', '').strip()
            html_lines.append(f'<p style="margin: 10px 0; font-size: 1.1em;"><span class="section-title">整体表现：</span>{content}</p>')
        
        elif line.startswith('**总得分**：'):
            content = line.replace('**总得分**：', '').strip()
            html_lines.append(f'<p style="margin: 10px 0; font-size: 1.2em;"><span class="section-title">总得分：</span><span style="color: #e74c3c; font-weight: bold; font-size: 1.1em;">{content}</span></p>')
        
        elif line.startswith('**统计信息**：'):
            html_lines.append('<p class="section-title">统计信息：</p>')
        
        elif line.startswith('**主要错误类型**：'):
            html_lines.append('<p class="section-title">主要错误类型：</p>')
        
        elif line.startswith('**改进建议**：'):
            html_lines.append('<p class="section-title">改进建议：</p>')
        
        # 列表项
        elif re.match(r'^\d+\.\s', line) or (line.startswith('- ') and '✓' not in line and '✗' not in line):
            content = re.sub(r'^\d+\.\s|^-\s', '', line)
            html_lines.append(f'<p style="margin: 3px 0 3px 30px;">• {content}</p>')
        
        else:
            # 其他内容
            html_lines.append(f'<p style="margin: 5px 0;">{line}</p>')
    
    # 组合HTML内容
    content = '\n'.join(html_lines)
    return html_template.format(content=content)